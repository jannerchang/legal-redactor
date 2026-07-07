from __future__ import annotations

import asyncio
import html
import json
import os
import re
import secrets
import subprocess
import urllib.error
import urllib.parse
import urllib.request
import base64
import importlib.util
import tempfile
from dataclasses import dataclass, replace
from datetime import datetime
from io import BytesIO
from pathlib import Path, PurePosixPath
from typing import Any
from zipfile import BadZipFile

from .config import PipelineConfig
from .cases import (
    CaseError,
    InvalidDiscordThreadError,
    InvalidWorkflowInputError,
    case_dir,
    case_root_from_source_dir,
    case_workflow_public,
    case_workflow_state,
    default_case_root,
    load_manifest,
    manifest_fields_for_case_dir,
    parse_discord_thread_id,
    persist_case_redaction,
    raise_for_forged_workflow_fields,
    record_hermes_thread_request,
    suggest_case_location_from_filenames as case_suggest_case_location_from_filenames,
    case_thread_binding_status,
    validate_case_folder_name,
    workflow_state_message,
)
from .counters import CN_ORDINALS, TypeCounters
from .io import is_encrypted_map, load_redaction_map_encrypted, redaction_map_from_json, redaction_map_to_json
from .local_config import config_value, load_json_config
from .models import MappingEntry, RedactedDocument, RedactionMap, sort_mapping_entries
from .org_masking import derived_organization_alias_cores
from .llm import is_noise_entity_text
from .pipeline import RedactionPipeline, _filter_noise_entity_mappings
from .restore import preview_restore, restore_docx
from .status import build_status_payload, ensure_mlx_server_ready
from .web_templates import (
    MAPPING_REVIEW_CATEGORY_LABELS,
    RESTORE_RISK_REASON_LABELS,
    _page,
    render_batch_redaction_result_page,
    render_home_page,
    render_redaction_result_page,
    render_status_panel,
)

try:
    from fastapi import FastAPI, File, Form, Request, UploadFile
    from fastapi.responses import HTMLResponse, JSONResponse
except ImportError as exc:
    raise RuntimeError("启动 Web UI 需要先安装依赖：pip install -r requirements.txt") from exc


app = FastAPI(title="本地法律文书脱敏系统", version="0.1.0")


@dataclass(frozen=True)
class InputDocument:
    source_file: str
    text: str


class DiscordApiError(RuntimeError):
    code = "discord_api_error"



SAMPLE_SUMMARY_KEYS = (
    "lookup_entries",
    "delete_blacklist_candidates",
    "suppressed_risky_entries",
    "manual_corrections",
    "false_positive_deletes",
    "missing_adds",
    "restore_unresolved_placeholders",
    "newest_sample_provenance",
    "regression_suggestions",
)

SUPPORTED_UPLOAD_SUFFIXES = {".txt", ".md", ".doc", ".docx", ".pdf"}


def _entity_group_is_noise(group: dict) -> bool:
    full = str(group.get("full_name", "")).strip()
    if full and is_noise_entity_text(full):
        return True
    aliases = group.get("aliases", [])
    if isinstance(aliases, list):
        for alias in aliases:
            alias_text = str(alias).strip()
            if alias_text and is_noise_entity_text(alias_text):
                return True
    return False


def _sanitize_redaction_map(redaction_map: RedactionMap) -> RedactionMap:
    filtered = _filter_noise_entity_mappings(redaction_map.mappings)
    if len(filtered) == len(redaction_map.mappings):
        return redaction_map
    return replace(redaction_map, mappings=sort_mapping_entries(filtered))


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok", "bind_host": "127.0.0.1", "network": "offline"}


@app.get("/api/status")
def api_status() -> dict:
    return _status_payload()


@app.post("/api/ensure-mlx")
def api_ensure_mlx() -> dict:
    item = ensure_mlx_server_ready()
    payload = item.to_dict()
    payload["status"] = "ok" if item.state in {"ready", "skipped"} else "error"
    return payload


def _mlx_not_ready_page(item) -> str | None:
    if item.state in {"ready", "skipped"}:
        return None
    return _page(
        "MLX 本地模型未就绪",
        f"<p>{html.escape(item.message)}</p>"
        f"<p><b>建议：</b>{html.escape(item.action)}</p>"
        "<p>也可双击桌面「启动文书脱敏系统」重新启动全套服务。</p>",
    )


def _redaction_failure_body(exc: Exception, *, enable_hanlp: bool) -> str:
    if enable_hanlp:
        suggestion = "建议：先取消勾选 HanLP，并确认 MLX 状态为就绪后重试。"
    else:
        suggestion = "建议：确认 MLX 状态为就绪后重试；当前未启用 HanLP，问题不在 HanLP 勾选项。"
    return f"<p>{html.escape(str(exc))}</p><p>{html.escape(suggestion)}</p>"


def _reject_forged_workflow_fields(body: dict) -> JSONResponse | None:
    try:
        raise_for_forged_workflow_fields(body)
    except InvalidWorkflowInputError as exc:
        return JSONResponse(
            {
                "status": "error",
                "code": exc.code,
                "fields": exc.fields,
                "message": str(exc),
            },
            status_code=400,
        )
    return None


def _reject_forged_workflow_form_data(form: dict) -> HTMLResponse | None:
    try:
        raise_for_forged_workflow_fields(form)
    except InvalidWorkflowInputError as exc:
        return HTMLResponse(
            _page(
                "请求无效",
                (
                    f'<p class="error">INVALID_INPUT：请求包含不能由浏览器提交的工作流决策字段。</p>'
                    f'<p class="hint">字段：{html.escape(", ".join(exc.fields))}</p>'
                ),
            ),
            status_code=400,
        )
    return None


def _case_error_response(message: str, *, code: str = "case_error", status_code: int = 400) -> JSONResponse:
    return JSONResponse(
        {
            "status": "error",
            "workflow_state": "attach_failed",
            "code": code,
            "message": _safe_public_error_message(message),
        },
        status_code=status_code,
    )


def _waiting_hermes_response() -> JSONResponse:
    return JSONResponse(
        {
            "status": "pending",
            "workflow_state": "waiting_hermes",
            "message": "等待 Hermes 写回 Discord 帖子链接",
        },
        status_code=202,
    )


@app.post("/api/save-to-local")
async def save_to_local(request: Request) -> JSONResponse:
    try:
        body = await request.json()
        directory = body.get("directory", "").strip()
        files = body.get("files", [])

        if not directory:
            return JSONResponse({"status": "error", "message": "保存目录不能为空"}, status_code=400)

        expanded_dir = os.path.abspath(os.path.expanduser(directory))

        try:
            os.makedirs(expanded_dir, exist_ok=True)
        except Exception as e:
            return JSONResponse({"status": "error", "message": f"创建/访问目录失败: {str(e)}"}, status_code=400)

        saved_paths = []
        for file_item in files:
            filename = file_item.get("filename", "").strip()
            content = file_item.get("content", "")
            if not filename:
                continue
            file_path = os.path.join(expanded_dir, filename)
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)
            saved_paths.append(file_path)

        if not saved_paths:
            return JSONResponse({"status": "error", "message": "没有需要保存的文件内容"}, status_code=400)

        return JSONResponse({
            "status": "success",
            "message": f"已成功保存 {len(saved_paths)} 个文件至本地目录：\n{expanded_dir}",
            "directory": expanded_dir,
            "saved_paths": saved_paths
        })
    except Exception as exc:
        return JSONResponse({"status": "error", "message": f"保存失败: {str(exc)}"}, status_code=500)


@app.post("/api/suggest-case-location")
async def suggest_case_location(request: Request) -> JSONResponse:
    body = await request.json()
    invalid = _reject_forged_workflow_fields(body)
    if invalid is not None:
        return invalid
    filenames = body.get("filenames", [])
    roots = []
    case_root = str(body.get("case_root", "")).strip()
    if case_root and not _is_default_case_root_value(case_root):
        roots.append(Path(case_root).expanduser())
    relative_paths = body.get("relative_paths") or body.get("upload_relative_paths") or []
    if _safe_upload_relative_paths(relative_paths):
        relative_suggestion = _suggest_case_location_from_relative_paths(
            relative_paths,
            roots or None,
            discord_thread_url=str(body.get("discord_thread_url", "")).strip(),
        )
        if relative_suggestion.get("status") != "not_found":
            return JSONResponse(relative_suggestion)
    suggestion = _suggest_case_location_from_filenames(
        filenames,
        roots or None,
        source_dir=str(body.get("source_dir") or body.get("upload_source_dir") or "").strip(),
        discord_thread_url=str(body.get("discord_thread_url", "")).strip(),
    )
    return JSONResponse(suggestion)


@app.post("/api/discord/send-redacted")
async def send_redacted_to_discord(request: Request) -> JSONResponse:
    body = await request.json()
    invalid = _reject_forged_workflow_fields(body)
    if invalid is not None:
        return invalid
    thread_url = str(body.get("discord_thread_url", "")).strip()
    filename = Path(str(body.get("filename", "redacted.txt"))).name or "redacted.txt"
    content = str(body.get("content", ""))
    if not thread_url:
        return _case_error_response("缺少 Discord 帖子链接")
    if not content:
        return _case_error_response("没有可发送的脱敏内容")
    try:
        thread_id = parse_discord_thread_id(thread_url)
    except InvalidDiscordThreadError as exc:
        return _case_error_response(str(exc), code=exc.code)
    try:
        result = _post_discord_thread_file(
            thread_id,
            filename,
            content,
            _safe_discord_attachment_message(filename, str(body.get("message", ""))),
        )
    except DiscordApiError as exc:
        return _case_error_response(str(exc), code=exc.code)
    except RuntimeError as exc:
        return _case_error_response(str(exc))
    return JSONResponse({"status": "success", "workflow_state": "sent_discord", **result})


@app.post("/api/discord/create-thread")
async def create_discord_thread(request: Request) -> JSONResponse:
    body = await request.json()
    invalid = _reject_forged_workflow_fields(body)
    if invalid is not None:
        return invalid
    case_folder = str(body.get("case_folder", "")).strip()
    case_cause = str(body.get("case_cause", "")).strip()
    if not case_folder:
        return _case_error_response("缺少案件文件夹名")
    source_dir = str(body.get("source_dir", "")).strip() or None
    try:
        source_root = case_root_from_source_dir(source_dir, case_folder) if source_dir else None
        case_root = str(source_root or str(body.get("case_root", "")).strip() or default_case_root())
        case_path = case_dir(case_root, case_folder)
        manifest = load_manifest(case_path) if (case_path / "manifest.json").exists() else None
    except CaseError as exc:
        return _case_error_response(str(exc), code=getattr(exc, "code", "case_error"))
    except Exception as exc:
        return _case_error_response(f"案件 manifest 读取失败: {exc}")
    if manifest and manifest.discord_thread_url:
        return JSONResponse(
            {
                "status": "bound",
                "workflow_state": "bound_thread",
                "thread_url": manifest.discord_thread_url,
                "thread_id": manifest.discord_thread_id,
                "message": "案件已绑定 Discord 帖子",
            }
        )
    if manifest and manifest.hermes_request_id:
        return JSONResponse(
            {
                "status": "pending",
                "workflow_state": "waiting_hermes",
                "request_id": manifest.hermes_request_id,
                "command_message_id": manifest.hermes_command_message_id,
                "channel_id": manifest.hermes_command_channel_id,
                "message": "已有 Hermes 建帖请求，继续等待写回帖子链接",
            }
        )
    request_id = str(body.get("request_id") or _new_discord_request_id())
    try:
        command = _case_creation_command(
            case_folder,
            request_id,
            case_cause,
        )
        result = _post_discord_channel_message(_discord_command_channel_id(), command)
        manifest = record_hermes_thread_request(
            case_root,
            case_folder,
            request_id,
            source_dir=source_dir,
            command_message_id=result.get("message_id", ""),
            command_channel_id=result.get("channel_id", ""),
        )
    except CaseError as exc:
        return _case_error_response(str(exc), code=getattr(exc, "code", "case_error"))
    except DiscordApiError as exc:
        return _case_error_response(str(exc), code=exc.code)
    except RuntimeError as exc:
        return _case_error_response(str(exc))
    return JSONResponse({
        "status": "pending",
        "workflow_state": "waiting_hermes",
        "request_id": manifest.hermes_request_id or request_id,
        "command_message_id": manifest.hermes_command_message_id or result.get("message_id", ""),
        "channel_id": manifest.hermes_command_channel_id or result.get("channel_id", ""),
        "message": "已发送建帖请求，等待 Hermes 通过 MCP 写回帖子链接",
    })


@app.post("/api/discord/attach-bound-thread")
async def attach_to_bound_discord_thread(request: Request) -> JSONResponse:
    body = await request.json()
    invalid = _reject_forged_workflow_fields(body)
    if invalid is not None:
        return invalid
    case_folder = str(body.get("case_folder", "")).strip()
    source_dir = str(body.get("source_dir", "")).strip() or None
    source_root = case_root_from_source_dir(source_dir, case_folder) if case_folder else None
    case_root = str(source_root or str(body.get("case_root", "")).strip() or default_case_root())
    filename = Path(str(body.get("filename", "redacted.txt"))).name or "redacted.txt"
    content = str(body.get("content", ""))
    map_json = str(body.get("map_json", ""))
    if not case_folder:
        return _case_error_response("缺少案件文件夹名")
    if not content:
        return _case_error_response("没有可发送的脱敏内容")
    try:
        case_path = case_dir(case_root, case_folder)
    except CaseError as exc:
        return _case_error_response(str(exc), code=getattr(exc, "code", "case_error"))
    try:
        manifest = load_manifest(case_path)
    except FileNotFoundError:
        return _waiting_hermes_response()
    except Exception as exc:
        return _case_error_response(f"案件 manifest 读取失败: {exc}")
    if not manifest.discord_thread_url:
        return _waiting_hermes_response()
    try:
        redaction_map = redaction_map_from_json(map_json)
    except Exception as exc:
        return _case_error_response(f"映射表解析失败: {exc}")
    try:
        thread_id = parse_discord_thread_id(manifest.discord_thread_url)
        result = _post_discord_thread_file(
            thread_id,
            filename=filename,
            content=content,
            message=_safe_discord_attachment_message(filename, str(body.get("message", ""))),
        )
        persist_case_redaction(
            case_root,
            case_folder,
            manifest.discord_thread_url,
            [RedactedDocument(source_file=filename, original_text="", redacted_text=content)],
            redaction_map,
            source_dir=source_dir,
        )
    except DiscordApiError as exc:
        return _case_error_response(str(exc), code=exc.code)
    except (CaseError, InvalidDiscordThreadError, RuntimeError) as exc:
        return _case_error_response(str(exc), code=getattr(exc, "code", "case_error"))
    return JSONResponse({
        "status": "success",
        "workflow_state": "sent_discord",
        "thread_url": manifest.discord_thread_url,
        "thread_id": thread_id,
        "case_folder": case_folder,
        **result,
    })


@app.post("/api/mapping/suggest-entry")
async def suggest_mapping_entry(request: Request) -> JSONResponse:
    body = await request.json()
    selected_text = str(body.get("selected_text", "")).strip()
    entity_type = str(body.get("entity_type", "")).strip()
    map_json = str(body.get("map_json", ""))
    if not selected_text:
        return JSONResponse({"status": "error", "message": "未选择文字"}, status_code=400)
    if len(selected_text) > 80:
        return JSONResponse({"status": "error", "message": "选择文字过长，请只选择一个实体"}, status_code=400)
    if entity_type not in {"person", "organization", "location"}:
        return JSONResponse({"status": "error", "message": "不支持的实体类型"}, status_code=400)
    try:
        redaction_map = redaction_map_from_json(map_json)
    except Exception as exc:
        return JSONResponse({"status": "error", "message": f"映射表解析失败: {exc}"}, status_code=400)

    existing = _find_mapping_by_original(redaction_map.mappings, selected_text)
    if existing:
        return JSONResponse({
            "status": "exists",
            "entry": existing.to_dict(),
            "message": f"已存在映射：{existing.original} → {existing.masked}",
        })

    entry = _suggest_manual_mapping_entry(selected_text, entity_type, redaction_map.mappings)
    return JSONResponse({"status": "success", "entry": entry.to_dict()})


@app.get("/", response_class=HTMLResponse)
def index() -> str:
    sample_info = ""
    status_panel = _render_status_panel(_status_payload())

    hanlp_attr = _hanlp_checked_attr()
    default_root_str = str(default_case_root())
    return render_home_page(status_panel, sample_info, hanlp_attr, default_root_str)

def _hanlp_checked_attr() -> str:
    # LLM 主路径下 HanLP 为可选增强；默认不勾选，避免与 MLX 同时占满内存导致进程被系统杀掉。
    return ""


def _status_payload() -> dict:
    return build_status_payload(mlx_timeout=0.4)


_render_status_panel = render_status_panel



def _is_default_case_root_value(value: str) -> bool:
    candidate = str(value or "").strip()
    if not candidate:
        return False
    try:
        return Path(candidate).expanduser().resolve() == default_case_root().expanduser().resolve()
    except OSError:
        return Path(candidate).expanduser() == default_case_root().expanduser()


@app.post("/analyze", response_class=HTMLResponse)
async def analyze_page(
    text: str = Form(default=""),
    llm_mode: str = Form(default="max-effect"),
    enable_llm: str | None = Form(default=None),
    file: UploadFile | None = File(default=None),
    files: list[UploadFile] = File(default=[]),
) -> str:
    try:
        documents = await _read_input_documents(text, file, files)
    except ValueError as exc:
        return _page("上传失败", str(exc))

    profile = "standard"
    llm_mode = "max-effect"
    mlx_block = _mlx_not_ready_page(ensure_mlx_server_ready())
    if mlx_block is not None:
        return mlx_block
    config = PipelineConfig.from_llm_mode(llm_mode, profile_name=profile)
    pipeline = RedactionPipeline(config=config)

    # 执行语义审计（后台线程，避免阻塞 /health 等轻量请求）
    raw_text = "\n\n".join(doc.text for doc in documents)
    analysis = await asyncio.to_thread(pipeline.analyze, raw_text)

    return _render_audit_dashboard(
        analysis=analysis,
        original_documents=documents,
        profile=profile,
        llm_mode=llm_mode
    )

def _render_audit_dashboard(
    analysis: dict,
    original_documents: list[InputDocument],
    profile: str,
    llm_mode: str,
    round_num: int = 0,
    previous_map_json: str = "{}",
    previous_deselected_json: str = "[]",
    locked_entries: list[MappingEntry] | None = None,
) -> str:
    locked_entries = locked_entries or []

    # 已锁定的实体展示
    locked_html = ""
    if locked_entries and round_num > 0:
        locked_rows = ""
        for e in locked_entries:
            locked_rows += f'<tr class="locked-row"><td><span class="tag tag-locked">已替换</span></td><td colspan="2">{html.escape(e.original)} → {html.escape(e.masked)}</td></tr>'
        locked_html = f"""
        <details class="locked-section" {'open' if len(locked_entries) <= 8 else ''}>
          <summary>已确认并替换 <span class="badge">{len(locked_entries)}</span> 条</summary>
          <table class="locked-table"><tbody>{locked_rows}</tbody></table>
        </details>
        """

    groups = analysis.get("entity_groups", [])
    groups_html = ""
    for g in groups:
        aliases = [a for a in g.get("aliases", []) if a]
        aliases_str = "、".join(aliases) if aliases else "无"
        entity_type_label = "公司/机构" if g.get("type") == "organization" else "个人"
        groups_html += f"""
        <tr class="entity-row">
          <td><input type="checkbox" checked name="group_{g.get('id')}_enabled" value="1"></td>
          <td><span class="tag type-{g.get('type')}">{entity_type_label}</span></td>
          <td><span class="tag role-{g.get('role')}">{g.get('role', '') or ''}</span></td>
          <td class="full-name"><strong>{html.escape(g.get('full_name', ''))}</strong></td>
          <td class="aliases">{html.escape(aliases_str)}</td>
          <td><input type="text" name="group_{g.get('id')}_mask" placeholder="自动生成" class="mask-input"></td>
        </tr>
        """

    locations = analysis.get("locations", [])
    locations_html = "".join(
        f'<li><label><input type="checkbox" checked name="loc_{idx}" value="{html.escape(str(l))}"> {html.escape(str(l))}</label></li>'
        for idx, l in enumerate(locations)
    )

    bundle_json = json.dumps([{"source_file": d.source_file, "text": d.text} for d in original_documents], ensure_ascii=False)

    round_badge = f' <span class="round-badge">第 {round_num + 1} 轮</span>' if round_num > 0 else ""
    subtitle = "（基于已脱敏文本的二次扫描）" if round_num > 0 else "（基于原文首次扫描）"

    return _page(
        f"分级确认 - 语义审计{round_badge}",
        f"""
        <nav><a href="/">返回首页</a></nav>
        <section class="info-card">
          <h2>识别到的主体与关联关系 {round_badge}</h2>
          <p class="hint">{subtitle} 大模型已自动将"全称"与"简称"归组。您可以取消勾选不需脱敏的项，或手动指定脱敏后的代号。</p>
          {locked_html}
          <form action="/redact/confirmed" method="post">
            <input type="hidden" name="profile" value="{profile}">
            <input type="hidden" name="llm_mode" value="{llm_mode}">
            <input type="hidden" name="bundle_json" value="{html.escape(bundle_json)}">
            <input type="hidden" name="analysis_json" value="{html.escape(json.dumps(analysis, ensure_ascii=False))}">
            <input type="hidden" name="round" value="{round_num}">
            <input type="hidden" name="previous_map_json" value="{html.escape(previous_map_json)}">
            <input type="hidden" name="previous_deselected_json" value="{html.escape(previous_deselected_json)}">
            {'<p class="hint" style="color:var(--muted)">未发现新实体 — 点击"完成脱敏"查看最终结果。</p>' if not groups and not locations else ''}
            {'<table class="audit-table">'
             '<thead><tr><th>脱敏</th><th>类型</th><th>角色</th><th>全称</th><th>关联简称</th><th>指定代号</th></tr></thead>'
             '<tbody>' + groups_html + '</tbody></table>' if groups else ''}
            {'<h3>识别到的地名</h3><ul class="tag-list">' + locations_html + '</ul>' if locations else ''}
            <div style="margin-top:30px; border-top:1px solid #eee; padding-top:20px; display:flex; gap:10px">
              <button type="submit" name="action" value="continue" class="btn">确认并继续分析</button>
              <button type="submit" name="action" value="finish" class="btn btn-secondary">完成脱敏</button>
            </div>
          </form>
        </section>
        <style>
          .audit-table {{ width:100%; border-collapse:collapse; margin:20px 0; }}
          .audit-table th, .audit-table td {{ padding:12px; border-bottom:1px solid #eee; text-align:left; }}
          .tag {{ padding:3px 8px; border-radius:4px; font-size:12px; font-weight:bold; }}
          .type-organization {{ background:#f0f7ff; color:#0052cc; }}
          .type-person {{ background:#fff7e6; color:#d46b08; }}
          .role-原告 {{ color:#52c41a; }}
          .role-被告 {{ color:#ff4d4f; }}
          .mask-input {{ border:1px solid #ddd; padding:5px; border-radius:4px; width:100px; }}
          .tag-list {{ list-style:none; padding:0; display:flex; flex-wrap:wrap; gap:10px; }}
          .tag-list li {{ background:#f5f5f5; padding:5px 12px; border-radius:20px; font-size:13px; }}
          .round-badge {{ font-size:13px; background:var(--accent); color:#fff; padding:2px 10px; border-radius:99px; font-weight:500 }}
          .locked-section {{ margin-bottom:18px; padding:12px; background:var(--bg); border-radius:var(--radius-sm); border:1px solid var(--border) }}
          .locked-section summary {{ cursor:pointer; font-weight:600; color:var(--muted); font-size:13px }}
          .locked-table {{ width:100%; border-collapse:collapse; margin-top:8px; font-size:12px }}
          .locked-table td {{ padding:6px 8px; border-bottom:1px solid var(--border) }}
          .locked-row {{ opacity:0.55 }}
          .tag-locked {{ background:#d4edda; color:#155724 }}
          .badge {{ background:var(--ink); color:#fff; padding:1px 7px; border-radius:99px; font-size:11px; font-weight:600 }}
        </style>
        """
    )


@app.post("/redact/confirmed", response_class=HTMLResponse)
async def redact_confirmed_page(request: Request) -> str:
    """增量脱敏：每轮确认后立即替换，再用已脱敏文本做下一轮分析。

    用户勾选的实体本轮生效，未勾选的不会出现在下一轮。
    每轮替换后的文本会传给 LLM 做二次审计，只展示新发现的实体。
    """
    form = await request.form()
    bundle_json = form.get("bundle_json", "")
    analysis_json = form.get("analysis_json", "")
    profile = "standard"
    llm_mode = "max-effect"
    round_num = int(form.get("round", "0"))
    previous_map_json = form.get("previous_map_json", "{}")
    previous_deselected_json = form.get("previous_deselected_json", "[]")
    action = form.get("action", "continue")

    docs_data = json.loads(bundle_json)
    is_batch = len(docs_data) > 1

    analysis = json.loads(analysis_json)
    analysis["entity_groups"] = [
        group
        for group in analysis.get("entity_groups", [])
        if isinstance(group, dict) and not _entity_group_is_noise(group)
    ]
    prev_data = json.loads(previous_map_json) if previous_map_json else {}
    prev_mappings_dicts: list[dict] = prev_data.get("mappings", [])
    prev_confirmed_texts: set[str] = set(prev_data.get("confirmed_texts", []))
    all_deselected_texts: set[str] = set(json.loads(previous_deselected_json))

    # 收集本轮勾选的实体文本和自定义掩码
    confirmed_texts: set[str] = set()
    user_masks: dict[str, str] = {}

    for g in analysis.get("entity_groups", []):
        gid = g.get("id")
        if form.get(f"group_{gid}_enabled"):
            full = g.get("full_name", "").strip()
            if full:
                confirmed_texts.add(full)
            for alias in g.get("aliases", []):
                alias = alias.strip()
                if alias:
                    confirmed_texts.add(alias)
            custom = (form.get(f"group_{gid}_mask") or "").strip()
            if custom and full:
                user_masks[full] = custom

    for idx, loc in enumerate(analysis.get("locations", [])):
        if form.get(f"loc_{idx}"):
            loc = str(loc).strip()
            if loc:
                confirmed_texts.add(loc)

    # 记录本轮未勾选的实体 → 永久排除
    for g in analysis.get("entity_groups", []):
        gid = g.get("id")
        if not form.get(f"group_{gid}_enabled"):
            full = g.get("full_name", "").strip()
            if full:
                all_deselected_texts.add(full)
            for alias in g.get("aliases", []):
                alias = alias.strip()
                if alias:
                    all_deselected_texts.add(alias)

    for idx, loc in enumerate(analysis.get("locations", [])):
        if not form.get(f"loc_{idx}"):
            loc = str(loc).strip()
            if loc:
                all_deselected_texts.add(loc)

    # 合并所有轮次的已确认实体
    all_confirmed = prev_confirmed_texts | confirmed_texts

    # 生成本轮新增的映射条目（按实体组：同组全称和简称共用掩码）
    counters = TypeCounters()
    for g in analysis.get("entity_groups", []):
        gid = g.get("id")
        if not form.get(f"group_{gid}_enabled"):
            continue
        full = g.get("full_name", "").strip()
        if not full:
            continue
        custom = (form.get(f"group_{gid}_mask") or "").strip()
        group_mask = custom or _simple_mask(full, counters)
        prev_mappings_dicts.append({
            "type": "manual", "original": full, "masked": group_mask,
            "role": g.get("role"), "source": "user_confirmed",
            "confidence": 1.0, "restore_by_default": True,
        })
        for alias in g.get("aliases", []):
            alias = alias.strip()
            if alias and alias != full:
                prev_mappings_dicts.append({
                    "type": "manual", "original": alias, "masked": group_mask,
                    "role": g.get("role"), "source": "user_confirmed_alias",
                    "confidence": 1.0, "restore_by_default": True,
                })

    # 地名使用独立掩码（强制用地点掩码，即使无后缀）
    for idx, loc in enumerate(analysis.get("locations", [])):
        if not form.get(f"loc_{idx}"):
            continue
        loc = str(loc).strip()
        if not loc:
            continue
        # 地名可能无后缀（如"石家庄""沧州"），不能用 _simple_mask（会误判为姓名）
        loc_masked = _simple_mask(loc, counters)
        if "自然人" in loc_masked:
            loc_masked = _guess_location_mask(loc)
        prev_mappings_dicts.append({
            "type": "manual", "original": loc, "masked": loc_masked,
            "role": None, "source": "user_confirmed",
            "confidence": 1.0, "restore_by_default": True,
        })

    # 合并去重（同一原文只保留一条）
    merged: dict[str, dict] = {}
    for m in prev_mappings_dicts:
        merged.setdefault(m["original"], m)
    all_mapping_dicts = list(merged.values())

    all_mappings = [
        MappingEntry(
            type=m.get("type", "manual"),
            original=m.get("original", ""),
            masked=m.get("masked", ""),
            role=m.get("role"),
            source=m.get("source", "user_confirmed"),
            confidence=float(m.get("confidence", 1.0)),
            restore_by_default=m.get("restore_by_default", True),
        )
        for m in all_mapping_dicts
    ]

    # 应用所有已确认映射
    pipeline = RedactionPipeline(config=PipelineConfig.offline_without_llm(profile))
    redaction_map = _sanitize_redaction_map(
        RedactionMap.create(mappings=all_mappings, mode=profile)
    )

    # 当前轮次的映射数据（传给下一轮）
    current_map_json = json.dumps({
        "mappings": all_mapping_dicts,
        "confirmed_texts": list(all_confirmed),
    }, ensure_ascii=False)
    deselected_json = json.dumps(list(all_deselected_texts), ensure_ascii=False)

    # ── 批量文档：直接展示最终结果（不支持多轮） ──
    if is_batch:
        redacted_docs: list[RedactedDocument] = []
        all_leaks: list = []
        for d in docs_data:
            rt = pipeline.apply_redaction_map(d["text"], redaction_map)
            lks = pipeline.scan_high_risk_leaks(rt)
            redacted_docs.append(RedactedDocument(
                source_file=d.get("source_file", ""),
                original_text=d["text"],
                redacted_text=rt,
                leaks=lks,
            ))
            all_leaks.extend(lks)
        return _render_batch_redaction_result(
            title="脱敏完成",
            documents=redacted_docs,
            redaction_map=redaction_map,
            review_candidates=[],
            leaks=all_leaks,
            warnings=[],
        )

    # ── 单文档：增量多轮确认 ──
    original_text = docs_data[0]["text"]
    redacted_text = pipeline.apply_redaction_map(original_text, redaction_map)

    # 点了"完成" → 直接展示最终结果
    if action == "finish":
        leaks = pipeline.scan_high_risk_leaks(redacted_text)
        return _render_redaction_result(
            title="脱敏完成",
            original_text=original_text,
            redacted_text=redacted_text,
            redaction_map=redaction_map,
            review_candidates=[],
            leaks=leaks,
            warnings=[],
        )

    # 对已脱敏文本做新一轮 LLM 分析
    config = PipelineConfig.from_llm_mode(llm_mode, profile_name=profile)
    pipeline2 = RedactionPipeline(config=config)
    new_analysis = await asyncio.to_thread(pipeline2.analyze, redacted_text)

    # 过滤掉已确认和已排除的实体
    new_groups = []
    for g in new_analysis.get("entity_groups", []):
        full = g.get("full_name", "").strip()
        aliases = [a.strip() for a in g.get("aliases", []) if a.strip()]
        all_texts = {full, *aliases}
        if all(t in all_confirmed or t in all_deselected_texts for t in all_texts if t):
            continue
        g["aliases"] = [a for a in aliases if a not in all_confirmed and a not in all_deselected_texts]
        new_groups.append(g)

    new_locations = [
        l for l in new_analysis.get("locations", [])
        if l not in all_confirmed and l not in all_deselected_texts
    ]

    new_analysis["entity_groups"] = new_groups
    new_analysis["locations"] = new_locations

    if new_groups or new_locations:
        return _render_audit_dashboard(
            analysis=new_analysis,
            original_documents=[InputDocument(source_file="", text=redacted_text)],
            profile=profile,
            llm_mode=llm_mode,
            round_num=round_num + 1,
            previous_map_json=current_map_json,
            previous_deselected_json=deselected_json,
            locked_entries=all_mappings,
        )
    else:
        leaks = pipeline.scan_high_risk_leaks(redacted_text)
        return _render_redaction_result(
            title="脱敏完成",
            original_text=original_text,
            redacted_text=redacted_text,
            redaction_map=redaction_map,
            review_candidates=[],
            leaks=leaks,
            warnings=[],
        )


@app.post("/redact", response_class=HTMLResponse)
async def redact_page(
    request: Request,
    text: str = Form(default=""),
    llm_mode: str = Form(default="max-effect"),
    enable_llm: str | None = Form(default=None),
    enable_hanlp: str | None = Form(default=None),
    hanlp_model: str = Form(default=""),
    enable_samples: str | None = Form(default=None),
    base_map_json: str = Form(default=""),
    case_folder: str = Form(default=""),
    discord_thread_url: str = Form(default=""),
    case_root: str = Form(default=""),
    upload_source_dir: str = Form(default=""),
    upload_relative_paths: str = Form(default=""),
    base_map_file: UploadFile | None = File(default=None),
    file: UploadFile | None = File(default=None),
    files: list[UploadFile] = File(default=[]),
    case_folder_files: list[UploadFile] = File(default=[]),
) -> str:
    form_invalid = _reject_forged_workflow_form_data(await request.form())
    if form_invalid is not None:
        return form_invalid

    try:
        documents = await _read_input_documents(text, file, files, case_folder_files)
    except ValueError as exc:
        return _page("上传失败", str(exc))

    base_redaction_map = None
    if base_map_json.strip() or (base_map_file and base_map_file.filename):
        try:
            map_text = await _read_restore_map_text(base_map_json, base_map_file)
            base_redaction_map = redaction_map_from_json(map_text)
        except Exception as exc:
            return _page("已有映射表解析失败", f"解析错误: {exc}")

    llm_mode = "max-effect"
    mlx_block = _mlx_not_ready_page(ensure_mlx_server_ready())
    if mlx_block is not None:
        return mlx_block
    config = PipelineConfig.from_llm_mode(
        llm_mode,
        profile_name="standard",
    )
    config = replace(
        config,
        enable_sample_library=bool(enable_samples),
        enable_hanlp_ner=bool(enable_hanlp),
        hanlp_model=hanlp_model.strip() or "MSRA_NER_ELECTRA_SMALL_ZH",
    )
    pipeline = RedactionPipeline(config=config)
    source_files = [item.source_file for item in documents]
    inferred_case_location = _resolve_case_location(upload_source_dir, source_files, upload_relative_paths)
    inferred_source_dir = str(inferred_case_location.get("matched_dir") or "")
    manual_case_root = "" if _is_default_case_root_value(case_root) else case_root.strip()
    effective_case_folder = case_folder.strip() or str(inferred_case_location.get("case_folder") or "")
    effective_case_root = manual_case_root or str(inferred_case_location.get("case_root") or "") or case_root.strip()
    effective_discord_thread_url = discord_thread_url.strip() or str(inferred_case_location.get("discord_thread_url") or "")
    try:
        if len(documents) > 1:
            result = await asyncio.to_thread(
                pipeline.redact_many,
                [(item.source_file, item.text) for item in documents],
                base_redaction_map=base_redaction_map,
            )
        else:
            result = await asyncio.to_thread(
                pipeline.redact,
                documents[0].text,
                source_file=documents[0].source_file,
                base_redaction_map=base_redaction_map,
            )
    except Exception as exc:
        return _page(
            "脱敏失败",
            _redaction_failure_body(exc, enable_hanlp=bool(enable_hanlp)),
        )
    result = replace(result, redaction_map=_sanitize_redaction_map(result.redaction_map))
    warnings = list(result.warnings)
    if len(documents) > 1:
        try:
            _persist_optional_case_redaction(
                effective_case_root,
                effective_case_folder,
                effective_discord_thread_url,
                result.documents,
                result.redaction_map,
                source_dir=inferred_source_dir,
            )
        except CaseError as exc:
            return _page("案件保存失败", f"保存错误: {exc}")
        except Exception as exc:
            return _page("案件保存失败", f"保存错误: {exc}")
        if effective_case_folder and effective_discord_thread_url:
            warnings.append(f"已保存到案件库：{effective_case_folder}")
        return _render_batch_redaction_result(
            "脱敏结果",
            result.documents,
            result.redaction_map,
            result.review_candidates,
            result.leaks,
            warnings,
            save_dir=inferred_source_dir,
            discord_thread_url=effective_discord_thread_url,
            case_root=effective_case_root,
            case_folder=effective_case_folder,
            source_dir=inferred_source_dir,
        )
    redacted_doc = RedactedDocument(
        source_file=documents[0].source_file,
        original_text=result.original_text,
        redacted_text=result.redacted_text,
        leaks=result.leaks,
    )
    try:
        _persist_optional_case_redaction(
            effective_case_root,
            effective_case_folder,
            effective_discord_thread_url,
            [redacted_doc],
            result.redaction_map,
            source_dir=inferred_source_dir,
        )
    except CaseError as exc:
        return _page("案件保存失败", f"保存错误: {exc}")
    except Exception as exc:
        return _page("案件保存失败", f"保存错误: {exc}")
    if effective_case_folder and effective_discord_thread_url:
        warnings.append(f"已保存到案件库：{effective_case_folder}")
    return _render_redaction_result(
        "脱敏结果",
        result.original_text,
        result.redacted_text,
        result.redaction_map,
        result.review_candidates,
        result.leaks,
        warnings,
        save_dir=inferred_source_dir,
        discord_thread_url=effective_discord_thread_url,
        case_root=effective_case_root,
        case_folder=effective_case_folder,
        source_dir=inferred_source_dir,
    )


@app.post("/redact/apply-map", response_class=HTMLResponse)
async def apply_map_page(
    original_text: str = Form(...),
    map_json: str = Form(...),
    original_bundle_json: str = Form(default=""),
) -> str:
    try:
        redaction_map = _sanitize_redaction_map(redaction_map_from_json(map_json))
    except Exception as exc:
        return _page("映射表解析失败", f"错误详情: {exc}")

    pipeline = RedactionPipeline(config=PipelineConfig(redaction_profile=RedactionProfile.from_preset("standard")))
    if original_bundle_json.strip():
        documents = _documents_from_bundle_json(original_bundle_json)
        redacted_documents = []
        all_leaks = []
        for doc in documents:
            redacted_text = pipeline.apply_redaction_map(doc.text, redaction_map)
            leaks = pipeline.scan_high_risk_leaks(redacted_text)
            redacted_documents.append(
                RedactedDocument(
                    source_file=doc.source_file,
                    original_text=doc.text,
                    redacted_text=redacted_text,
                    leaks=leaks,
                )
            )
            all_leaks.extend(leaks)
        return _render_batch_redaction_result(
            title="应用映射表结果",
            documents=redacted_documents,
            redaction_map=redaction_map,
            review_candidates=[],
            leaks=all_leaks,
            warnings=["已重新应用您上传/修改后的映射表。"],
        )

    redacted_text = pipeline.apply_redaction_map(original_text, redaction_map)
    leaks = pipeline.scan_high_risk_leaks(redacted_text)
    return _render_redaction_result(
        title="应用映射表结果",
        original_text=original_text,
        redacted_text=redacted_text,
        redaction_map=redaction_map,
        review_candidates=[],
        leaks=leaks,
        warnings=["已重新应用您上传/修改后的映射表。"],
    )


@app.post("/redact/apply-edited-map", response_class=HTMLResponse)
async def apply_edited_map_page(request: Request) -> str:
    form = await request.form()
    form_invalid = _reject_forged_workflow_form_data(form)
    if form_invalid is not None:
        return form_invalid

    original_text = form.get("original_text", "")
    original_bundle_json = form.get("original_bundle_json", "")
    save_dir = str(form.get("save_dir", ""))
    discord_thread_url = str(form.get("discord_thread_url", ""))
    case_root = str(form.get("case_root", ""))
    case_folder = str(form.get("case_folder", ""))
    source_dir = str(form.get("source_dir", ""))
    map_version = form.get("map_version", "1.0")
    map_created_at = form.get("map_created_at", "")
    map_mode = form.get("map_mode", "normal")
    map_source_file = form.get("map_source_file", "")

    map_type = form.getlist("map_type")
    map_original = form.getlist("map_original")
    map_masked = form.getlist("map_masked")
    map_role = form.getlist("map_role")
    map_source = form.getlist("map_source")
    map_confidence = form.getlist("map_confidence")
    map_reason = form.getlist("map_reason")
    map_restore_by_default = form.getlist("map_restore_by_default")
    row_delete = form.getlist("row_delete")
    remap_placeholders = str(form.get("remap_placeholders", "")).strip() == "1"

    redaction_map = _sanitize_redaction_map(
        _redaction_map_from_rows(
            version=map_version, created_at=map_created_at, mode=map_mode,
            source_file=map_source_file, map_type=map_type, map_original=map_original,
            map_masked=map_masked, map_role=map_role, map_source=map_source,
            map_confidence=map_confidence, map_reason=map_reason,
            map_restore_by_default=map_restore_by_default,
            row_delete=row_delete,
        )
    )
    warnings = ["已手动调整映射表。"]
    if remap_placeholders:
        redaction_map = replace(
            redaction_map,
            mappings=sort_mapping_entries(_renumber_mapping_placeholders(redaction_map.mappings)),
        )
        warnings.append("已按当前保留的映射重新排列占位符。")
    pipeline = RedactionPipeline(config=PipelineConfig.offline_without_llm())
    documents = _documents_from_bundle_json(original_bundle_json)
    if documents:
        redacted_documents = _apply_map_to_documents(pipeline, documents, redaction_map)
        leaks = [lk for d in redacted_documents for lk in d.leaks]
        return _render_batch_redaction_result(
            "编辑映射后结果",
            redacted_documents,
            redaction_map,
            [],
            leaks,
            warnings,
            save_dir=save_dir,
            discord_thread_url=discord_thread_url,
            case_root=case_root,
            case_folder=case_folder,
            source_dir=source_dir,
        )
    redacted_text = pipeline.apply_redaction_map(original_text, redaction_map)
    leaks = pipeline.scan_high_risk_leaks(redacted_text)
    return _render_redaction_result(
        "编辑映射后结果",
        original_text,
        redacted_text,
        redaction_map,
        [],
        leaks,
        warnings,
        save_dir=save_dir,
        discord_thread_url=discord_thread_url,
        case_root=case_root,
        case_folder=case_folder,
        source_dir=source_dir,
    )


def _source_indicates_manual(source: str) -> bool:
    normalized = (source or "").strip().lower()
    return normalized.startswith(("manual", "user", "selection"))


def _source_indicates_sample(source: str) -> bool:
    return "sample" in (source or "").strip().lower()


def _review_candidate_text_set(review_candidates: list) -> set[str]:
    values: set[str] = set()
    for candidate in review_candidates or []:
        text = getattr(candidate, "text", None)
        if text:
            values.add(str(text))
    return values


def _classify_mapping_review_row(
    entry: MappingEntry,
    *,
    original_entry: dict[str, Any] | MappingEntry | None = None,
    deleted: bool = False,
    review_candidate_texts: set[str] | None = None,
    is_new_row: bool = False,
) -> list[str]:
    categories: list[str] = []
    review_candidate_texts = review_candidate_texts or set()
    if entry.confidence < 0.85 or entry.original in review_candidate_texts:
        categories.append("low_confidence")
    if _source_indicates_manual(entry.source) or (
        is_new_row and entry.source not in {"rule", "regex", "llm"}
    ):
        categories.append("manual_added")
    if original_entry is not None:
        old_masked = original_entry.masked if isinstance(original_entry, MappingEntry) else str(original_entry.get("masked", ""))
        if old_masked and old_masked != entry.masked:
            categories.append("modified")
    if deleted:
        categories.append("delete_candidate")
    if _restore_risk_reasons(entry, deleted=deleted):
        categories.append("restore_risk")
    if _source_indicates_sample(entry.source):
        categories.append("sample_reused")
    return [name for name in MAPPING_REVIEW_CATEGORY_LABELS if name in categories]


def _restore_risk_reasons(entry: MappingEntry, *, deleted: bool = False) -> list[dict[str, str]]:
    reasons: list[dict[str, str]] = []
    if deleted:
        reasons.append({
            "reason_code": "delete_candidate",
            "message": RESTORE_RISK_REASON_LABELS["delete_candidate"],
        })
    if not entry.masked:
        reasons.append({
            "reason_code": "empty_mask",
            "message": RESTORE_RISK_REASON_LABELS["empty_mask"],
        })
    return reasons


def _sample_entry_original(entry: dict[str, Any]) -> str:
    if entry.get("action") == "modify":
        return str(entry.get("new_original") or entry.get("old_original") or "")
    return str(entry.get("original") or "")


def _sample_entry_core(entry: dict[str, Any]) -> dict[str, str]:
    keys = (
        "action",
        "type",
        "original",
        "masked",
        "old_original",
        "new_original",
        "old_masked",
        "new_masked",
        "reason",
    )
    return {key: str(entry.get(key) or "") for key in keys if entry.get(key) not in (None, "")}


def _sample_effective_delta(existing_entries: list[dict[str, Any]], incoming_entries: list[dict[str, Any]]) -> dict[str, int]:
    existing_by_original = {
        _sample_entry_original(entry): _sample_entry_core(entry)
        for entry in existing_entries
        if _sample_entry_original(entry)
    }
    delta = {"created": 0, "updated": 0, "unchanged": 0}
    for entry in incoming_entries:
        original = _sample_entry_original(entry)
        if not original:
            continue
        current_core = _sample_entry_core(entry)
        previous_core = existing_by_original.get(original)
        if previous_core is None:
            delta["created"] += 1
        elif previous_core == current_core:
            delta["unchanged"] += 1
        else:
            delta["updated"] += 1
        existing_by_original[original] = current_core
    return delta


def _load_sample_entries_for_delta(samples_dir: str | Path | None = None) -> list[dict[str, Any]]:
    from ._samples import _auto_sample_path

    path = _auto_sample_path(samples_dir) if samples_dir is not None else _auto_sample_path()
    if not path.exists():
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []
    entries = data.get("entries", [])
    return entries if isinstance(entries, list) else []


def _summary_item_from_entry(entry: dict[str, Any], *, action: str | None = None) -> dict[str, Any]:
    actual_action = action or str(entry.get("action") or "")
    if actual_action == "modify":
        original = str(entry.get("new_original") or entry.get("old_original") or "")
        masked = str(entry.get("new_masked") or entry.get("old_masked") or "")
    else:
        original = str(entry.get("original") or "")
        masked = str(entry.get("masked") or "")
    item = {
        "action": actual_action,
        "type": str(entry.get("type") or "other"),
        "original": original,
        "masked": masked,
    }
    reason = str(entry.get("reason") or "").strip()
    if reason:
        item["reason"] = reason
    return item


def _empty_sample_summary(source_file: str = "") -> dict[str, Any]:
    return {
        "lookup_entries": [],
        "delete_blacklist_candidates": [],
        "suppressed_risky_entries": [],
        "manual_corrections": 0,
        "false_positive_deletes": 0,
        "missing_adds": 0,
        "restore_unresolved_placeholders": None,
        "newest_sample_provenance": {
            "source": "web_ui",
            "source_file_present": bool(source_file),
        },
        "regression_suggestions": [],
    }


def _sample_provenance(source_file: str = "", sample_path: str | Path | None = None) -> dict[str, Any]:
    provenance: dict[str, Any] = {
        "source": "web_ui",
        "source_file_present": bool(source_file),
    }
    if sample_path:
        path = Path(sample_path)
        provenance["sample_file"] = path.name
        try:
            provenance["sample_updated_at"] = datetime.fromtimestamp(path.stat().st_mtime).isoformat()
        except OSError:
            provenance["sample_updated_at"] = None
    return provenance


def _build_sample_save_summary(
    entries: list[dict[str, Any]],
    *,
    skipped_risky_entries: list[dict[str, Any]] | None = None,
    source_file: str = "",
    sample_path: str | Path | None = None,
) -> dict[str, Any]:
    from ._samples import is_sample_lookup_allowed

    skipped_risky_entries = skipped_risky_entries or []
    summary = _empty_sample_summary(source_file)
    summary["newest_sample_provenance"] = _sample_provenance(source_file, sample_path)
    for entry in entries:
        action = str(entry.get("action") or "")
        if action in {"add", "modify"}:
            item = _summary_item_from_entry(entry, action=action)
            lookup_entry = {
                "action": action,
                "type": item["type"],
                "original": item["original"],
                "masked": item["masked"],
            }
            if is_sample_lookup_allowed(entry, item["original"], item["masked"]):
                summary["lookup_entries"].append(lookup_entry)
            else:
                summary["suppressed_risky_entries"].append({
                    **lookup_entry,
                    "reason_code": "lookup_guard",
                    "message": RESTORE_RISK_REASON_LABELS["lookup_guard"],
                })
            if action == "add":
                summary["missing_adds"] += 1
        elif action == "delete":
            item = _summary_item_from_entry(entry, action=action)
            summary["delete_blacklist_candidates"].append({
                "action": action,
                "type": item["type"],
                "original": item["original"],
                "reason_code": "delete_candidate",
                "message": RESTORE_RISK_REASON_LABELS["delete_candidate"],
            })
            summary["false_positive_deletes"] += 1
    summary["suppressed_risky_entries"].extend(skipped_risky_entries)
    summary["manual_corrections"] = len(entries) + len(skipped_risky_entries)
    suggestions: list[str] = []
    if entries or skipped_risky_entries:
        suggestions.append(".venv/bin/python -m pytest tests/test_sample_integration.py")
    if any(item.get("action") in {"add", "modify"} for item in entries):
        suggestions.append(".venv/bin/python -m pytest tests/test_web_app.py")
    summary["regression_suggestions"] = suggestions
    return {key: summary[key] for key in SAMPLE_SUMMARY_KEYS}


def _sample_summary_response(msg: str, summary: dict[str, Any], *, cls: str = "") -> HTMLResponse:
    payload: dict[str, Any] = {
        "type": "sample_summary",
        "msg": msg,
        "summary": summary,
    }
    if cls:
        payload["cls"] = cls
    return HTMLResponse(
        f"<script>parent.postMessage({json.dumps(payload, ensure_ascii=False)},\"*\")</script>"
    )


@app.post("/redact/save-sample", response_class=HTMLResponse)
async def save_sample_page(request: Request) -> str:
    form = await request.form()
    map_type = form.getlist("map_type")
    map_original = form.getlist("map_original")
    map_masked = form.getlist("map_masked")
    map_role = form.getlist("map_role")
    map_source = form.getlist("map_source")
    map_confidence = form.getlist("map_confidence")
    map_reason = form.getlist("map_reason")
    map_restore_by_default = form.getlist("map_restore_by_default")
    row_delete = form.getlist("row_delete")
    map_source_file = form.get("map_source_file", "")
    original_mapping_json = form.get("original_mapping_json", "")

    from . import _samples as samples_module
    from ._samples import is_global_delete_sample_allowed, save_sample_auto

    try:
        original_data = json.loads(original_mapping_json) if original_mapping_json else {}
    except json.JSONDecodeError:
        original_data = {}
    original_mappings = original_data.get("mappings", [])
    original_index = {e.get("original", ""): e for e in original_mappings}

    deleted = set(str(r) for r in row_delete)
    edited_index: dict[str, str] = {}
    edited_types: dict[str, str] = {}
    edited_reasons: dict[str, str] = {}
    for i in range(max(len(map_original), len(map_masked))):
        if str(i) in deleted: continue
        orig = (map_original[i] if i < len(map_original) else "").strip()
        masked = (map_masked[i] if i < len(map_masked) else "").strip()
        t = (map_type[i] if i < len(map_type) else "other").strip()
        reason = (map_reason[i] if i < len(map_reason) else "").strip()
        if orig and masked:
            edited_index[orig] = masked
            edited_types[orig] = t
            if reason:
                edited_reasons[orig] = reason

    entries: list[dict] = []
    processed: set[str] = set()
    skipped_risky_deletes: list[dict[str, Any]] = []
    for i_str in deleted:
        try:
            i = int(i_str)
            if i < len(map_original):
                orig = map_original[i].strip()
                if orig and orig not in processed:
                    entry = {"action": "delete", "type": map_type[i] if i < len(map_type) else "other", "original": orig}
                    reason = (map_reason[i] if i < len(map_reason) else "").strip()
                    if reason:
                        entry["reason"] = reason
                    if is_global_delete_sample_allowed(entry):
                        entries.append(entry)
                    else:
                        skipped_risky_deletes.append({
                            "action": "delete",
                            "type": entry["type"],
                            "original": orig,
                            "reason_code": "risky_delete_guard",
                            "message": "短中文人名未写入全局黑名单",
                        })
                    processed.add(orig)
        except (ValueError, IndexError):
            continue
    for orig, masked in edited_index.items():
        if orig in processed: continue
        processed.add(orig)
        t = edited_types.get(orig, "other")
        reason = edited_reasons.get(orig, "")
        if orig in original_index:
            old_masked = original_index[orig].get("masked", "")
            if masked != old_masked:
                entry = {"action": "modify", "type": t, "old_original": orig, "new_original": orig, "old_masked": old_masked, "new_masked": masked}
                if reason:
                    entry["reason"] = reason
                entries.append(entry)
            # keep 条目不保存（识别正确的无需记录）
        else:
            entry = {"action": "add", "type": t, "original": orig, "masked": masked}
            if reason:
                entry["reason"] = reason
            entries.append(entry)

    if not entries:
        summary = _build_sample_save_summary(
            entries,
            skipped_risky_entries=skipped_risky_deletes,
            source_file=str(map_source_file),
        )
        if skipped_risky_deletes:
            return _sample_summary_response(
                "短中文人名未写入全局黑名单，请用修改映射或规则修正处理",
                summary,
                cls="warn",
            )
        return _sample_summary_response("无变化，未追加", summary)

    existing_sample_entries = _load_sample_entries_for_delta(samples_module.DEFAULT_SAMPLES_DIR)
    effective_delta = _sample_effective_delta(existing_sample_entries, entries)

    try:
        sample_path = save_sample_auto(entries, source=map_source_file or "web_ui")
    except Exception as exc:
        return HTMLResponse(f'<script>parent.postMessage({{type:"toast",msg:"保存失败:{html.escape(str(exc))}",cls:"warn"}},"*")</script>')

    new_count = sum(1 for e in entries if e["action"] in ("add", "modify"))
    del_count = sum(1 for e in entries if e["action"] == "delete")
    summary = _build_sample_save_summary(
        entries,
        skipped_risky_entries=skipped_risky_deletes,
        source_file=str(map_source_file),
        sample_path=sample_path,
    )
    msg = (
        f'已处理 {len(entries)} 条 | 新增 {effective_delta["created"]} | '
        f'更新 {effective_delta["updated"]} | 未变化 {effective_delta["unchanged"]} | '
        f'匹配 {new_count} | 黑名单 {del_count}'
    )
    if skipped_risky_deletes:
        msg += f' | 跳过短人名黑名单 {len(skipped_risky_deletes)}'
    return _sample_summary_response(msg, summary, cls="warn" if skipped_risky_deletes else "")


def _diagnose_sample_entry(entry: dict) -> str:
    action = entry.get("action", "")
    orig = entry.get("original") or entry.get("new_original", "")
    masked = entry.get("masked") or entry.get("new_masked", "")
    manual_reason = str(entry.get("reason") or "").strip()
    if manual_reason:
        return html.escape(manual_reason)

    if action == "delete":
        matched_rules = []
        # 1. 手机/电话
        if re.search(r"(?<!\d)(?:\+?86[-\s]?)?1[3-9]\d{9}(?!\d)", orig):
            matched_rules.append("手机号正则")
        # 2. 身份证
        if re.search(r"(?<![0-9Xx])\d{6}(?:18|19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx](?![0-9Xx])", orig):
            matched_rules.append("身份证号正则")
        # 3. 信用代码
        if re.search(r"(?<![A-Z0-9])[0-9A-Z]{18}(?![A-Z0-9])", orig):
            matched_rules.append("信用代码正则")
        # 4. 案号
        if re.search(r"[（(][12]\d{3}[）)][\u4e00-\u9fa5A-Za-z0-9]{1,16}?(?:知民初|知民终|执异|执复|民辖终|民辖初|民辖|民初|民终|民申|民再|行初|行终|行申|刑初|刑终|刑申|刑再|商初|商终|破申|执|民撤|民特|民保|强清|管辖)", orig):
            matched_rules.append("案号结构化正则")
        # 5. 地名/行政区划
        if re.search(r"[\u4e00-\u9fa5]{2,6}?(?:省|自治区|市|自治州|盟|区|县|旗|镇|乡|街道|村|社区)$", orig):
            matched_rules.append("启发式地名匹配")
        # 6. 常见人名兜底
        if len(orig) in (2, 3):
            matched_rules.append("姓名兜底匹配")
        # 7. 机构/公司
        if any(orig.endswith(sfx) for sfx in ["有限责任公司", "股份有限公司", "集团有限公司", "有限公司", "公司", "集团", "律师事务所", "会计师事务所", "经营部", "商行", "工作室", "厂", "店"]):
            matched_rules.append("机构后缀特征")

        if not matched_rules:
            matched_rules.append("LLM语义审计或规则兜底")

        rules_str = "、".join(matched_rules)
        return f"<span style='color:var(--danger);font-weight:500'>误匹配为实体</span>（触发「{html.escape(rules_str)}」）。<b>已加入黑名单，下次分析相同文本将自动豁免，不再误判！</b>"

    elif action == "modify":
        old_masked = entry.get("old_masked", "")
        new_masked = entry.get("new_masked", "")
        return f"<span style='color:#e65100;font-weight:500'>修正脱敏掩码</span>（从「{html.escape(old_masked)}」修正为「{html.escape(new_masked)}」）。<b>已载入精准映射，下次直接以此格式脱敏该词。</b>"

    elif action == "add":
        return f"<span style='color:#1565c0;font-weight:500'>手动新增实体</span>。<b>已载入精准映射，下次该词将直接脱敏为「{html.escape(masked)}」。</b>"

    elif action == "keep":
        return f"<span style='color:#2e7d32;font-weight:500'>确认无误（保留）</span>。<b>已记录，下次将直接命中并直接脱敏。</b>"

    return "已作为脱敏样本库记录并投入使用。"


@app.get("/samples/edit", response_class=HTMLResponse)
def edit_samples_page() -> str:
    from ._samples import _auto_sample_path
    filepath = _auto_sample_path()
    entries = []
    if filepath.exists():
        try:
            data = json.loads(filepath.read_text(encoding="utf-8"))
            entries = data.get("entries", [])
        except (json.JSONDecodeError, KeyError):
            pass

    rows = ""
    for i, e in enumerate(entries):
        action = e.get("action", "?")
        orig = e.get("original") or e.get("new_original", "")
        masked = e.get("masked") or e.get("new_masked", "")
        old = e.get("old_original", "")
        reason = e.get("reason", "")
        action_label = {"keep": "保留", "delete": "黑名单", "add": "新增", "modify": "修改"}.get(action, action)
        row_class = "style='opacity:.6'" if action == "delete" else ""
        row_diagnose = _diagnose_sample_entry(e)
        rows += f"""<tr {row_class}>
          <td><span class="tag tag-{action}">{action_label}</span></td>
          <td><input name="orig_{i}" value="{html.escape(orig)}" style="width:180px"></td>
          <td><input name="masked_{i}" value="{html.escape(masked)}" style="width:140px"></td>
          <td style="font-size:11px;color:var(--muted)">{html.escape(old)}</td>
          <td><textarea name="reason_{i}" rows="2" style="width:220px" placeholder="为什么删除/修改/添加">{html.escape(reason)}</textarea></td>
          <td style="font-size:12px;color:var(--ink);max-width:320px;word-break:break-all">{row_diagnose}</td>
          <td>
            <button class="btn btn-sm" onclick="saveRow({i},this)">保存</button>
            <a href="/samples/delete/{i}" class="btn btn-sm btn-secondary" style="margin-left:4px">删除</a>
          </td>
        </tr>"""
    rows += f"""<tr>
      <td><select id="new-action" style="width:80px"><option value="add">新增</option><option value="delete">黑名单</option></select></td>
      <td><input id="new-orig" placeholder="原文" style="width:180px"></td>
      <td><input id="new-masked" placeholder="替换为" style="width:140px"></td>
      <td></td>
      <td><textarea id="new-reason" rows="2" style="width:220px" placeholder="为什么新增/删除"></textarea></td>
      <td></td>
      <td><button class="btn btn-sm" onclick="saveNewRow({len(entries)},this)">＋</button></td>
    </tr>"""

    return _page(
        "编辑样本库",
        f"""
        <nav><a href="/">返回首页</a></nav>
        <style>
          .tag{{display:inline-block;padding:2px 8px;border-radius:99px;font-size:11px;font-weight:600}}
          .tag-keep{{background:#e8f5e9;color:#2e7d32}}
          .tag-delete{{background:#fce4ec;color:#c62828}}
          .tag-add{{background:#e3f2fd;color:#1565c0}}
          .tag-modify{{background:#fff3e0;color:#e65100}}
        </style>
        <section>
          <h2>样本库（{len(entries)} 条）</h2>
          <p class="hint">编辑后自动保存。删除操作立即生效。</p>
          <table>
            <thead><tr><th>类型</th><th>原文</th><th>替换为</th><th>旧值</th><th>修改理由</th><th>诊断与优化分析</th><th>操作</th></tr></thead>
            <tbody>{rows}</tbody>
          </table>
        </section>
        """,
    )


@app.post("/samples/update/{idx}")
async def update_sample_entry(idx: int, request: Request) -> JSONResponse:
    from ._samples import _auto_sample_path, save_sample_auto
    filepath = _auto_sample_path()
    body = await request.json()
    if not filepath.exists():
        return JSONResponse({"msg": "样本库为空"})
    data = json.loads(filepath.read_text(encoding="utf-8"))
    entries = data.get("entries", [])
    if 0 <= idx < len(entries):
        e = entries[idx]
        action = e.get("action", "")
        if action in ("keep", "add"):
            e["original"] = body.get("original", e.get("original", ""))
            e["masked"] = body.get("masked", e.get("masked", ""))
        elif action == "modify":
            e["new_original"] = body.get("original", e.get("new_original", ""))
            e["new_masked"] = body.get("masked", e.get("new_masked", ""))
        e["reason"] = body.get("reason", e.get("reason", "")).strip()
        save_sample_auto([e], source=e.get("source", "samples_edit"))
        return JSONResponse({"msg": "已保存"})
    return JSONResponse({"msg": "索引无效"}, status_code=400)


@app.post("/samples/add")
async def add_sample_entry(request: Request) -> JSONResponse:
    from ._samples import is_global_delete_sample_allowed, save_sample_auto
    body = await request.json()
    action = body.get("action", "add")
    orig = body.get("original", "").strip()
    masked = body.get("masked", "").strip()
    reason = body.get("reason", "").strip()
    if not orig:
        return JSONResponse({"msg": "原文不能为空"}, status_code=400)
    if action == "delete":
        entry = {"action": "delete", "type": "manual", "original": orig}
        if reason:
            entry["reason"] = reason
        if not is_global_delete_sample_allowed(entry):
            return JSONResponse({"msg": "短中文人名不写入全局黑名单，请改用精确映射校准。"}, status_code=400)
        save_sample_auto([entry], source="samples_edit")
    else:
        if not masked:
            return JSONResponse({"msg": "替换为不能为空"}, status_code=400)
        entry = {"action": "add", "type": "manual", "original": orig, "masked": masked}
        if reason:
            entry["reason"] = reason
        save_sample_auto([entry], source="samples_edit")
    return JSONResponse({"msg": "已添加"})


@app.get("/samples/delete/{idx}", response_class=HTMLResponse)
def delete_sample_entry(idx: int) -> str:
    from ._samples import _auto_sample_path
    filepath = _auto_sample_path()
    if not filepath.exists():
        return _page("错误", '<p>样本库为空。</p><a href="/samples/edit">返回</a>')
    try:
        data = json.loads(filepath.read_text(encoding="utf-8"))
        entries = data.get("entries", [])
        if 0 <= idx < len(entries):
            entries.pop(idx)
            data["entries"] = entries
            data["total"] = len(entries)
            filepath.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return _page("已删除", f'<p>第 {idx+1} 条已删除。</p><a href="/samples/edit">返回样本库</a>')
    except Exception as exc:
        return _page("错误", f'<p>{html.escape(str(exc))}</p><a href="/samples/edit">返回</a>')


@app.get("/samples/compact", response_class=HTMLResponse)
def compact_samples_page() -> str:
    from ._samples import compact_samples
    compact_samples()
    return _page("整理完成", '<p class="success">样本库已去重并优化。</p><nav><a href="/">返回首页</a></nav>')


def _render_case_workflow_panel(
    *,
    case_root: str = "",
    case_folder: str = "",
    discord_thread_url: str = "",
    saved_local: bool = False,
    hermes_requested: bool = False,
    attach_status: str = "",
    attach_error: str | None = None,
) -> str:
    status = case_workflow_public(
        case_root=case_root,
        case_folder=case_folder,
        discord_thread_url=discord_thread_url,
        saved_local=saved_local,
        hermes_requested=hermes_requested,
        attach_status=attach_status,
        attach_error=attach_error,
    )
    state = str(status.get("workflow_state", "not_saved"))
    message = str(status.get("message") or workflow_state_message(state, attach_error=attach_error))
    case_label = case_folder.strip() or "未选择案件"
    thread_url = str(status.get("discord_thread_url") or discord_thread_url or "").strip()
    next_action = {
        "not_saved": "可先保存到本地案件，或填写案件目录后请求 Hermes 建帖。",
        "saved_local": "可继续绑定 Discord 帖子。",
        "bound_thread": "可发送脱敏附件到 Discord。",
        "sent_discord": "等待 Discord/Hermes 后续审查起草。",
        "waiting_hermes": "稍后继续检查并绑定帖子。",
        "attach_failed": "检查失败原因后可再次发送附件。",
    }.get(state, "可继续处理案件流程。")
    thread_html = (
        f'<a href="{html.escape(thread_url, quote=True)}" target="_blank" rel="noopener">打开 Discord 帖子</a>'
        if thread_url
        else '<span class="hint">尚未绑定 Discord 帖子</span>'
    )
    manifest = status.get("manifest") if isinstance(status.get("manifest"), dict) else {}
    restore = manifest.get("restore") if isinstance(manifest.get("restore"), dict) else {}
    mapping_label = "已就绪" if manifest.get("mapping_present") else "缺失"
    restore_filename = str(restore.get("restored_filename") or "")
    restore_label = restore_filename or {
        "missing_map": "等待映射表",
        "no_restore_yet": "尚无还原文件",
        "metadata_unknown": "已有文件，缺少元数据",
        "restore_failed": "最近还原失败",
    }.get(str(restore.get("status") or ""), "尚无还原文件")
    unresolved = restore.get("unresolved_placeholder_count")
    unresolved_label = "未知" if unresolved is None else str(unresolved)
    return f"""
        <section class="case-workflow-panel" data-workflow-state="{html.escape(state, quote=True)}">
          <div class="workflow-head">
            <span class="workflow-pill workflow-{html.escape(state, quote=True)}">{html.escape(_workflow_state_label(state))}</span>
            <strong>案件流程状态</strong>
          </div>
          <div class="workflow-grid">
            <span><b>案件</b>{html.escape(case_label)}</span>
            <span><b>状态</b>{html.escape(message)}</span>
            <span><b>线程</b>{thread_html}</span>
            <span><b>下一步</b>{html.escape(next_action)}</span>
            <span><b>映射表</b>{html.escape(mapping_label)}</span>
            <span><b>还原状态</b>{html.escape(restore_label)}</span>
            <span><b>未解析占位符</b>{html.escape(unresolved_label)}</span>
          </div>
        </section>
    """


def _workflow_state_label(state: str) -> str:
    return {
        "not_saved": "未保存",
        "saved_local": "本地已保存",
        "bound_thread": "已绑定",
        "sent_discord": "已发送",
        "waiting_hermes": "等 Hermes",
        "attach_failed": "附件失败",
    }.get(state, state)


def _should_apply_auto_prefill(current_value: str, previous_auto_value: str) -> bool:
    current = current_value.strip()
    return not current or current == previous_auto_value


@app.post("/restore/preview", response_class=HTMLResponse)
async def restore_preview_page(
    text: str = Form(default=""),
    file: UploadFile | None = File(default=None),
    map_json: str = Form(default=""),
    map_file: UploadFile | None = File(default=None),
    restore_docx_format: str | None = Form(default=None),
) -> str:
    try:
        redacted_text = text.strip()
        redacted_docx_bytes: bytes | None = None
        redacted_filename = ""
        if file and file.filename:
            data = await file.read()
            redacted_filename = file.filename
            if _suffix_for_filename(file.filename) == ".docx":
                redacted_docx_bytes = data
                redacted_text = _docx_bytes_to_text(data)
            else:
                redacted_text = _decode_text_bytes(data, file.filename)
        map_text = await _read_restore_map_text(map_json, map_file)

        if not map_text or not redacted_text:
            return _page("参数缺失", '<nav><a href="/">返回</a></nav><section class="warning"><p>请粘贴或上传脱敏文本/Word，并提供映射表。</p></section>')

        redaction_map = redaction_map_from_json(map_text)
        if redacted_docx_bytes is not None and restore_docx_format:
            return _render_docx_restore_result(redacted_docx_bytes, redacted_filename, redaction_map)
        preview = preview_restore(redacted_text, redaction_map)
    except Exception as exc:
        return _page("还原错误", f'<nav><a href="/">返回</a></nav><section class="warning"><p>{html.escape(str(exc))}</p></section>')

    default_dir = os.path.expanduser("~/Desktop")
    restored_url = _data_download("restored.txt", "text/plain", preview.restored_text)
    restored_rows = "".join(
        f"<tr><td>{html.escape(i.type)}</td><td>{html.escape(i.masked)}</td><td>{html.escape(i.original)}</td></tr>"
        for i in preview.restored_entries
    )
    return _page("还原预览", f"""
        <nav><a href="/">返回首页</a></nav>
        <div class="downloads"><a download="restored.txt" href="{restored_url}" class="btn">下载还原文本</a></div>

        <section class="local-save-section" style="border-left: 4px solid var(--accent); background: linear-gradient(135deg, var(--surface) 0%, rgba(26, 122, 109, 0.02) 100%); padding: 18px 24px; border-radius: var(--radius); border: 1px solid var(--border); margin-bottom: 18px; box-shadow: var(--shadow);">
          <div style="display: flex; align-items: center; justify-content: space-between; flex-wrap: wrap; gap: 15px;">
            <div style="flex: 1; min-width: 280px;">
              <h3 style="margin: 0 0 6px 0; font-size: 14px; font-weight: 600; color: var(--ink); display: flex; align-items: center; gap: 6px;">
                <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="var(--accent)" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round" class="feather feather-folder"><path d="M22 19a2 2 0 0 1-2 2H4a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h5l2 3h9a2 2 0 0 1 2 2z"></path></svg>
                本地直接保存 <span class="hint" style="font-weight: normal; font-size: 11px; margin-left: 4px;">(保存至本地任意文件夹)</span>
              </h3>
              <div style="display: flex; gap: 8px; align-items: center; margin-top: 8px;">
                <span class="hint" style="white-space: nowrap; font-weight: 500;">保存路径:</span>
                <input type="text" id="local-save-dir" value="{html.escape(default_dir)}" style="flex: 1; min-width: 200px; padding: 6px 10px; border-radius: 6px; font-family: monospace; font-size: 13px;" placeholder="例如: ~/Desktop">
              </div>
            </div>
            <div style="display: flex; gap: 8px; align-items: center; flex-wrap: wrap; margin-top: 8px;">
              <button type="button" class="btn btn-sm" onclick="saveToLocalPath([{{filename: 'restored.txt', content: document.getElementById('restored-output').value}}], this)">保存还原文本</button>
            </div>
          </div>
          <script>
            (function(){{
              var savedDir = localStorage.getItem('last_local_save_dir');
              if (savedDir) {{
                var inp = document.getElementById('local-save-dir');
                if (inp) inp.value = savedDir;
              }}
            }})();
          </script>
        </section>

        <section class="grid">
          <div><h2>脱敏文本</h2><textarea rows="20" readonly>{html.escape(redacted_text)}</textarea></div>
          <div><h2>还原后</h2><textarea id="restored-output" rows="20" readonly>{html.escape(preview.restored_text)}</textarea></div>
        </section>
        <section>
          <h2>已还原</h2><table><thead><tr><th>类型</th><th>占位符</th><th>原文</th></tr></thead><tbody>{restored_rows}</tbody></table>
          <details><summary>差异预览</summary><pre>{html.escape(preview.diff)}</pre></details>
        </section>
    """)


def _render_docx_restore_result(
    redacted_docx_bytes: bytes,
    redacted_filename: str,
    redaction_map: RedactionMap,
) -> str:
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        input_path = temp_path / "redacted.docx"
        output_path = temp_path / "restored.docx"
        input_path.write_bytes(redacted_docx_bytes)
        replacements = restore_docx(input_path, output_path, redaction_map)
        restored_bytes = output_path.read_bytes()

    stem = Path(redacted_filename or "restored.docx").stem
    restored_filename = f"{stem}.restored.docx"
    restored_url = _binary_download(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        restored_bytes,
    )
    restored_rows = "".join(
        f"<tr><td>{html.escape(i.type)}</td><td>{html.escape(i.masked)}</td><td>{html.escape(i.original)}</td></tr>"
        for i in redaction_map.mappings
    )
    return _page("Word 还原完成", f"""
        <nav><a href="/">返回首页</a></nav>
        <div class="downloads">
          <a download="{html.escape(restored_filename)}" href="{restored_url}" class="btn" data-no-intercept="true">下载还原 Word</a>
        </div>
        <section class="info-card">
          <h2>还原完成</h2>
          <p>已按映射表生成保留格式的 Word 文档。</p>
          <p class="hint">替换次数：{replacements}；映射条目：{len(redaction_map.mappings)}。</p>
        </section>
        <section>
          <h2>映射表条目</h2>
          <table><thead><tr><th>类型</th><th>占位符</th><th>原文</th></tr></thead><tbody>{restored_rows}</tbody></table>
        </section>
    """)


def _render_redaction_result(
    title: str,
    original_text: str,
    redacted_text: str,
    redaction_map: RedactionMap,
    review_candidates: list,
    leaks: list,
    warnings: list[str],
    save_dir: str = "",
    discord_thread_url: str = "",
    case_root: str = "",
    case_folder: str = "",
    source_dir: str = "",
) -> str:
    default_dir = save_dir.strip() or os.path.expanduser("~/Desktop")
    map_json = redaction_map_to_json(redaction_map)
    from .debug_trace import debug_trace_from_parts, debug_trace_to_json

    debug_json = debug_trace_to_json(
        debug_trace_from_parts(
            mode=redaction_map.mode,
            source_file=redaction_map.source_file,
            mappings=redaction_map.mappings,
            documents=[
                {
                    "source_file": redaction_map.source_file,
                    "original_text": original_text,
                    "redacted_text": redacted_text,
                }
            ],
            review_candidates=review_candidates,
            leaks=leaks,
            warnings=warnings,
        )
    )
    leaks_html = "".join(f"<li><strong>{html.escape(lk.type)}</strong>: <mark>{html.escape(lk.text)}</mark></li>" for lk in leaks)
    warnings_html = "".join(f"<li>{html.escape(w)}</li>" for w in warnings)
    redacted_filename = "redacted.txt"
    redacted_filename_json = json.dumps(redacted_filename, ensure_ascii=False)
    redacted_url = _data_download(redacted_filename, "text/plain", redacted_text)
    map_url = _data_download("redaction_map.json", "application/json", map_json)
    debug_url = _data_download("debug_trace.json", "application/json", debug_json)
    discord_create_section = _discord_create_thread_section(
        discord_thread_url=discord_thread_url,
        case_root=case_root,
        case_folder=case_folder,
        source_dir=source_dir or save_dir,
        filename=redacted_filename,
        textarea_id="redacted-output",
        map_textarea_id="mapping-json-output",
        message_id="discord-create-message",
    )
    discord_section = _discord_send_section(discord_thread_url, redacted_filename, "redacted-output", "discord-message")
    workflow_panel = _render_case_workflow_panel(
        case_root=case_root,
        case_folder=case_folder,
        discord_thread_url=discord_thread_url,
        saved_local=bool(case_folder),
    )
    mapping_review_toolbar = _render_mapping_review_toolbar(redaction_map, review_candidates)
    sample_summary_panel = _render_sample_summary_panel()
    review_candidate_texts_json = _review_candidate_texts_json(review_candidates)
    review_html = "".join(
        "<tr><td>{}</td><td>{}</td><td>{}</td><td>{:.2f}</td><td>{}</td></tr>".format(
            html.escape(c.type), html.escape(c.text), html.escape(c.source),
            c.confidence, html.escape(c.reason or ""))
        for c in review_candidates
    )
    original_highlight = _highlight_replaced_text(original_text, redaction_map.mappings)
    redacted_highlight = _highlight_replaced_text(redacted_text, redaction_map.mappings, reverse=True)
    mapping_edit_rows = _render_mapping_edit_rows(redaction_map, review_candidates)
    return render_redaction_result_page(
        title=title,
        redacted_filename=redacted_filename,
        redacted_url=redacted_url,
        map_url=map_url,
        debug_url=debug_url,
        workflow_panel=workflow_panel,
        default_dir=default_dir,
        redacted_filename_json=redacted_filename_json,
        save_dir=save_dir,
        discord_create_section=discord_create_section,
        discord_section=discord_section,
        leaks_html=leaks_html,
        warnings_html=warnings_html,
        original_highlight=original_highlight,
        redacted_text=redacted_text,
        redacted_highlight=redacted_highlight,
        mapping_review_toolbar=mapping_review_toolbar,
        sample_summary_panel=sample_summary_panel,
        original_text=original_text,
        map_json=map_json,
        review_candidate_texts_json=review_candidate_texts_json,
        debug_json=debug_json,
        discord_thread_url=discord_thread_url,
        case_root=case_root,
        case_folder=case_folder,
        source_dir=source_dir,
        redaction_map=redaction_map,
        mapping_edit_rows=mapping_edit_rows,
        review_html=review_html,
    )


def _render_batch_redaction_result(
    title: str,
    documents: list[RedactedDocument],
    redaction_map: RedactionMap,
    review_candidates: list,
    leaks: list,
    warnings: list[str],
    save_dir: str = "",
    discord_thread_url: str = "",
    case_root: str = "",
    case_folder: str = "",
    source_dir: str = "",
) -> str:
    default_dir = save_dir.strip() or os.path.expanduser("~/Desktop")

    # 构建各个独立文件的脱敏文本列表供 JS 使用
    individual_files = []
    for index, document in enumerate(documents, start=1):
        out_name = f"document-{index}.redacted.txt"
        individual_files.append({"filename": out_name, "content": document.redacted_text})
    individual_files_json = json.dumps(individual_files, ensure_ascii=False)

    map_json = redaction_map_to_json(redaction_map)
    bundle_json = _documents_bundle_json(documents)
    combined_redacted = "\n\n".join(d.redacted_text for d in documents)
    from .debug_trace import debug_trace_from_parts, debug_trace_to_json

    debug_json = debug_trace_to_json(
        debug_trace_from_parts(
            mode=redaction_map.mode,
            source_file=redaction_map.source_file,
            mappings=redaction_map.mappings,
            documents=[
                {
                    "source_file": document.source_file,
                    "original_text": document.original_text,
                    "redacted_text": document.redacted_text,
                }
                for document in documents
            ],
            review_candidates=review_candidates,
            leaks=leaks,
            warnings=warnings,
        )
    )
    leaks_html = "".join(f"<li><strong>{html.escape(lk.type)}</strong>: <mark>{html.escape(lk.text)}</mark></li>" for lk in leaks)
    warnings_html = "".join(f"<li>{html.escape(w)}</li>" for w in warnings)
    map_url = _data_download("redaction_map.json", "application/json", map_json)
    debug_url = _data_download("debug_trace.json", "application/json", debug_json)
    combined_filename = "batch.redacted.txt"
    combined_filename_json = json.dumps(combined_filename, ensure_ascii=False)
    redacted_url = _data_download(combined_filename, "text/plain", combined_redacted)
    discord_create_section = _discord_create_thread_section(
        discord_thread_url=discord_thread_url,
        case_root=case_root,
        case_folder=case_folder,
        source_dir=source_dir or save_dir,
        filename=combined_filename,
        textarea_id="redacted-output",
        map_textarea_id="mapping-json-output",
        message_id="discord-create-message-batch",
    )
    discord_section = _discord_send_section(discord_thread_url, combined_filename, "redacted-output", "discord-message-batch")
    workflow_panel = _render_case_workflow_panel(
        case_root=case_root,
        case_folder=case_folder,
        discord_thread_url=discord_thread_url,
        saved_local=bool(case_folder),
    )
    mapping_review_toolbar = _render_mapping_review_toolbar(redaction_map, review_candidates)
    sample_summary_panel = _render_sample_summary_panel()
    review_candidate_texts_json = _review_candidate_texts_json(review_candidates)
    doc_sections = "".join(
        f'<article class="doc-result">'
        f'<h3>{html.escape(d.source_file)}</h3>'
        f'<h4>原文高亮</h4><div class="highlight-box original-highlight selection-add-source">{_highlight_replaced_text(d.original_text, redaction_map.mappings)}</div>'
        f'<h4>脱敏文</h4><div class="highlight-box redacted-highlight">{_highlight_replaced_text(d.redacted_text, redaction_map.mappings, reverse=True)}</div>'
        f'</article>'
        for d in documents
    )
    mapping_edit_rows = _render_mapping_edit_rows(redaction_map, review_candidates)
    return render_batch_redaction_result_page(
        title=title,
        combined_filename=combined_filename,
        redacted_url=redacted_url,
        map_url=map_url,
        debug_url=debug_url,
        combined_redacted=combined_redacted,
        workflow_panel=workflow_panel,
        default_dir=default_dir,
        combined_filename_json=combined_filename_json,
        save_dir=save_dir,
        individual_files_json=individual_files_json,
        discord_create_section=discord_create_section,
        discord_section=discord_section,
        leaks_html=leaks_html,
        warnings_html=warnings_html,
        doc_sections=doc_sections,
        mapping_review_toolbar=mapping_review_toolbar,
        sample_summary_panel=sample_summary_panel,
        bundle_json=bundle_json,
        map_json=map_json,
        review_candidate_texts_json=review_candidate_texts_json,
        debug_json=debug_json,
        discord_thread_url=discord_thread_url,
        case_root=case_root,
        case_folder=case_folder,
        source_dir=source_dir,
        redaction_map=redaction_map,
        mapping_edit_rows=mapping_edit_rows,
    )


def _discord_create_thread_section(
    *,
    discord_thread_url: str,
    case_root: str,
    case_folder: str,
    source_dir: str,
    filename: str,
    textarea_id: str,
    map_textarea_id: str,
    message_id: str,
) -> str:
    if discord_thread_url.strip() or not case_folder.strip():
        return ""
    status_id = f"{message_id}-status"
    link_id = f"{message_id}-link"
    cause_id = f"{message_id}-case-cause"
    return (
        f'<section class="local-save-section" style="border-left: 4px solid #5865f2; background: linear-gradient(135deg, var(--surface) 0%, rgba(88, 101, 242, 0.04) 100%); padding: 18px 24px; border-radius: var(--radius); border: 1px solid var(--border); margin-bottom: 18px; box-shadow: var(--shadow);">'
        f'<div style="display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:15px;">'
        f'<div style="flex:1;min-width:280px;">'
        f'<h3 style="margin:0 0 8px 0;font-size:14px;font-weight:600;color:var(--ink);">请求 Hermes 新建案件帖</h3>'
        f'<p class="hint" style="margin:0;">向 Discord 指令频道发送建帖请求；Hermes 建帖后通过 MCP 写回链接，系统随后发送脱敏附件并写入本地案件库：{html.escape(case_folder)}</p>'
        f'<textarea id="{html.escape(message_id, quote=True)}" rows="2" placeholder="建帖后发送附件时附言" style="margin-top:10px;max-width:680px;">脱敏文件已生成，请见附件。</textarea>'
        f'</div>'
        f'<div style="display:flex;flex-direction:column;gap:8px;align-items:flex-start;min-width:220px;">'
        f'<button type="button" class="btn discord-create-thread-button" '
        f'data-case-root="{html.escape(case_root, quote=True)}" '
        f'data-case-folder="{html.escape(case_folder, quote=True)}" '
        f'data-source-dir="{html.escape(source_dir, quote=True)}" '
        f'data-filename="{html.escape(filename, quote=True)}" '
        f'data-textarea-id="{html.escape(textarea_id, quote=True)}" '
        f'data-map-textarea-id="{html.escape(map_textarea_id, quote=True)}" '
        f'data-message-id="{html.escape(message_id, quote=True)}" '
        f'data-status-id="{html.escape(status_id, quote=True)}" '
        f'data-case-cause-id="{html.escape(cause_id, quote=True)}" '
        f'data-link-id="{html.escape(link_id, quote=True)}">'
        f'请求 Hermes 建帖并绑定</button>'
        f'<input type="text" id="{html.escape(cause_id, quote=True)}" placeholder="案由（目录只有案号时填写）" style="width:100%;max-width:260px;">'
        f'<span id="{html.escape(status_id, quote=True)}" class="hint"></span>'
        f'<a id="{html.escape(link_id, quote=True)}" href="#" target="_blank" style="display:none;font-size:12px;">打开 Discord 帖子</a>'
        f'</div></div></section>'
    )


def _discord_send_section(discord_thread_url: str, filename: str, textarea_id: str, message_id: str) -> str:
    thread_url = discord_thread_url.strip()
    if not thread_url:
        return ""
    default_message = "脱敏文件已生成，请见附件。"
    status_id = f"{message_id}-status"
    return (
        f'<section class="local-save-section" style="border-left: 4px solid #5865f2; background: linear-gradient(135deg, var(--surface) 0%, rgba(88, 101, 242, 0.04) 100%); padding: 18px 24px; border-radius: var(--radius); border: 1px solid var(--border); margin-bottom: 18px; box-shadow: var(--shadow);">'
        f'<div style="display: flex; align-items: flex-start; justify-content: space-between; flex-wrap: wrap; gap: 14px;">'
        f'<div style="flex: 1; min-width: 280px;">'
        f'<h3 style="margin: 0 0 8px 0; font-size: 14px; font-weight: 600; color: var(--ink);">发送到 Discord 帖子</h3>'
        f'<p class="hint" style="margin: 0 0 8px 0;">只发送脱敏文本附件，不发送映射表。</p>'
        f'<textarea id="{html.escape(message_id, quote=True)}" rows="3" style="width: 100%; min-height: 70px; resize: vertical;">{html.escape(default_message)}</textarea>'
        f'</div>'
        f'<div style="display: flex; align-items: flex-start; justify-content: flex-end; flex-direction: column; gap: 8px; min-height: 120px;">'
        f'<button type="button" class="btn discord-send-button" '
        f'data-thread-url="{html.escape(thread_url, quote=True)}" '
        f'data-filename="{html.escape(filename, quote=True)}" '
        f'data-textarea-id="{html.escape(textarea_id, quote=True)}" '
        f'data-message-id="{html.escape(message_id, quote=True)}" '
        f'data-status-id="{html.escape(status_id, quote=True)}">'
        f'一键发送到 Discord</button>'
        f'<span id="{html.escape(status_id, quote=True)}" class="hint" style="min-height: 18px;"></span>'
        f'</div></div></section>'
    )


def _persist_optional_case_redaction(
    case_root: str,
    case_folder: str,
    discord_thread_url: str,
    documents: list[RedactedDocument],
    redaction_map: RedactionMap,
    *,
    source_dir: str = "",
) -> None:
    has_case_folder = bool(case_folder.strip())
    has_thread_url = bool(discord_thread_url.strip())
    if not has_case_folder and not has_thread_url:
        return
    if not has_case_folder and has_thread_url:
        raise CaseError("填写 Discord 帖子链接时必须同时填写案件文件夹名")
    source_root = case_root_from_source_dir(source_dir, case_folder)
    root = str(source_root) if source_root is not None else case_root.strip()
    root = root or str(default_case_root())
    persist_case_redaction(
        root,
        case_folder,
        discord_thread_url,
        documents,
        redaction_map,
        source_dir=source_dir.strip() or None,
    )


def _resolve_case_location(upload_source_dir: str, source_files: list[str], upload_relative_paths: str = "") -> dict[str, object]:
    source_dir = upload_source_dir.strip()
    if source_dir:
        return _suggest_case_location_from_filenames(source_files, source_dir=source_dir)
    relative_suggestion = _suggest_case_location_from_relative_paths(upload_relative_paths)
    if relative_suggestion.get("status") == "ok":
        return relative_suggestion
    suggestion = _suggest_case_location_from_filenames(source_files)
    if suggestion.get("status") == "ok":
        return suggestion
    return {"status": "not_found"}


def _suggest_case_location_from_relative_paths(
    relative_paths: object,
    search_roots: list[Path] | None = None,
    *,
    discord_thread_url: str = "",
) -> dict[str, object]:
    paths = _safe_upload_relative_paths(relative_paths)
    case_folder = _case_folder_from_relative_paths(paths)
    if not case_folder:
        return {"status": "not_found", "workflow_state": "not_saved", "evidence": []}

    roots = search_roots or _case_location_search_roots()
    existing_dirs: list[Path] = []
    for root in roots:
        candidate = (Path(root).expanduser() / case_folder)
        if candidate.exists():
            existing_dirs.append(candidate.resolve())

    unique_dirs = sorted({path for path in existing_dirs}, key=str)
    if len(unique_dirs) > 1:
        return {
            "status": "ambiguous",
            "workflow_state": "not_saved",
            "confidence": 0.0,
            "matches": [str(path) for path in unique_dirs[:8]],
            "candidates": [_case_folder_hint_summary(path.parent, case_folder, matched_dir=path) for path in unique_dirs[:8]],
            "evidence": [{"kind": "ambiguous_case_directory", "count": len(unique_dirs)}],
        }

    if unique_dirs:
        result = _case_folder_hint_summary(unique_dirs[0].parent, case_folder, matched_dir=unique_dirs[0])
        result["confidence"] = 0.98
    else:
        root = Path(roots[0]).expanduser() if roots else default_case_root()
        result = _case_folder_hint_summary(root, case_folder)
        result["confidence"] = 0.86

    result["status"] = "ok"
    result["evidence"] = [
        {"kind": "upload_relative_path", "case_folder": case_folder},
        *list(result.get("evidence", [])),
    ]
    requested_thread = discord_thread_url.strip()
    if requested_thread:
        _apply_requested_thread_preflight(result, requested_thread)
    return result


def _safe_upload_relative_paths(value: object) -> list[str]:
    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return []
        try:
            parsed = json.loads(raw)
        except json.JSONDecodeError:
            parsed = [raw]
    else:
        parsed = value
    if not isinstance(parsed, list):
        return []

    paths: list[str] = []
    for item in parsed:
        path = str(item or "").replace("\\", "/").strip()
        if not path or path.startswith("/") or path.startswith("~"):
            continue
        pure = PurePosixPath(path)
        if any(part in {"", ".", ".."} for part in pure.parts):
            continue
        if pure.name.startswith("._") or _suffix_for_filename(pure.name) not in SUPPORTED_UPLOAD_SUFFIXES:
            continue
        paths.append(str(pure))
    return paths


def _case_folder_from_relative_paths(paths: list[str]) -> str:
    folder = ""
    for value in paths:
        parts = PurePosixPath(value).parts
        if len(parts) < 2:
            continue
        current = parts[0]
        if not folder:
            folder = current
        elif folder != current:
            return ""
    if not folder:
        return ""
    try:
        return validate_case_folder_name(folder)
    except CaseError:
        return ""


def _case_folder_hint_summary(case_root: Path, case_folder: str, *, matched_dir: Path | None = None) -> dict[str, object]:
    case_path = matched_dir or (case_root / case_folder)
    result: dict[str, object] = {
        "case_folder": case_folder,
        "case_root": str(Path(case_root).expanduser()),
        "matched_dir": str(matched_dir) if matched_dir else "",
        "ambiguous": False,
        "conflict": False,
    }
    evidence = list(result.get("evidence", []))
    manifest_data = manifest_fields_for_case_dir(case_path)
    result.update(manifest_data)
    result["evidence"] = evidence + list(manifest_data.get("evidence", []))
    result.setdefault("workflow_state", case_workflow_state(discord_thread_url=str(result.get("discord_thread_url", ""))))
    return result


def _apply_requested_thread_preflight(result: dict[str, object], requested_thread: str) -> None:
    try:
        binding = case_thread_binding_status(
            str(result.get("case_root", "")),
            str(result.get("case_folder", "")),
            requested_thread,
        )
    except CaseError as exc:
        result.update(
            {
                "status": "conflict",
                "conflict": True,
                "conflict_code": getattr(exc, "code", "case_error"),
                "conflict_message": str(exc),
            }
        )
        return
    if binding.get("conflict"):
        result.update(
            {
                "status": "conflict",
                "conflict": True,
                "conflict_code": binding.get("code"),
                "conflict_message": binding.get("message"),
            }
        )


def _discord_bot_token(config: dict | None = None) -> str:
    config = config if config is not None else load_json_config("LEGAL_REDACTOR_API_CONFIG", "api.local.json")
    token = os.environ.get("LEGAL_REDACTOR_DISCORD_BOT_TOKEN") or config_value(config, "discord_bot_token")
    if not token or token.startswith("optional-"):
        raise RuntimeError("未配置 Discord bot token")
    return str(token)


def _new_discord_request_id() -> str:
    return f"lr_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{secrets.token_hex(3)}"


def _case_creation_command(
    case_folder: str,
    request_id: str,
    case_cause: str = "",
    *,
    case_root: str = "",
    source_dir: str = "",
) -> str:
    _ = case_root, source_dir
    folder = validate_case_folder_name(case_folder)
    lines = [
        f"新建案件，{_case_creation_title(folder, case_cause)}",
        f"请求ID：{_safe_discord_request_id(request_id)}",
        f"案件目录：{folder}",
    ]
    return "\n".join(lines)


def _safe_discord_request_id(request_id: str) -> str:
    value = re.sub(r"[^A-Za-z0-9_.:-]+", "_", request_id.strip())[:80]
    return value or _new_discord_request_id()


def _case_creation_title(case_folder: str, case_cause: str = "") -> str:
    value = case_folder.strip()
    cause = _clean_case_cause(case_cause)
    paren_match = re.match(r"^[（(]\s*(\d{4})\s*[）)]\s*(\d{1,8})(?:\s*号)?\s*(.*)$", value)
    space_match = re.match(r"^(\d{4})\s+(\d{1,8})(?:\s*号)?\s*(.*)$", value)
    match = paren_match or space_match
    if not match:
        return f"{value} {cause}" if cause else value
    year, number, tail = match.groups()
    normalized = f"（{year}）{number}"
    tail = tail.strip() or cause
    return f"{normalized} {tail}" if tail else normalized


def _clean_case_cause(case_cause: str) -> str:
    value = re.sub(r"\s+", " ", case_cause).strip()
    value = re.sub(r"^案由\s*[:：]\s*", "", value)
    if _contains_local_path_text(value):
        return ""
    return value[:80]


def _contains_local_path_text(value: str) -> bool:
    return bool(
        re.search(r"(^|\s)(~?/|/Users/|/Volumes/|/private/|/var/folders/|[A-Za-z]:[\\/]|\\\\)", value)
        or re.search(r"[\\/].+[\\/]", value)
    )


def _discord_command_channel_id() -> str:
    config = load_json_config("LEGAL_REDACTOR_API_CONFIG", "api.local.json")
    return str(
        os.environ.get("LEGAL_REDACTOR_DISCORD_COMMAND_CHANNEL_ID")
        or config_value(config, "discord_command_channel_id")
        or "1501248343823880345"
    )


def _post_discord_channel_message(channel_id: str, content: str) -> dict[str, str]:
    config = load_json_config("LEGAL_REDACTOR_API_CONFIG", "api.local.json")
    token = _discord_bot_token(config)
    channel_id = str(channel_id).strip()
    if not channel_id:
        raise RuntimeError("未配置 Discord 指令频道 id")

    payload = {"content": content.strip()[:1900]}
    request = urllib.request.Request(
        f"https://discord.com/api/v10/channels/{channel_id}/messages",
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bot {token}",
            "Content-Type": "application/json",
            "User-Agent": "legal-redactor/0.1",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=30) as response:
            data = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        exc.read()
        raise DiscordApiError(f"Discord 指令发送失败: HTTP {exc.code}") from exc
    except OSError as exc:
        raise DiscordApiError("Discord 指令发送失败: 网络不可达") from exc
    return {
        "message_id": str(data.get("id", "")),
        "channel_id": str(data.get("channel_id") or channel_id),
    }


def _post_discord_thread_file(thread_id: str, filename: str, content: str, message: str = "") -> dict[str, str]:
    config = load_json_config("LEGAL_REDACTOR_API_CONFIG", "api.local.json")
    token = _discord_bot_token(config)
    message = _safe_discord_attachment_message(filename, message)

    payload = {
        "content": message,
        "attachments": [{"id": 0, "filename": filename}],
    }
    fields = [
        ("payload_json", "application/json", None, json.dumps(payload, ensure_ascii=False).encode("utf-8")),
        ("files[0]", "text/plain; charset=utf-8", filename, content.encode("utf-8")),
    ]
    body, content_type = _multipart_form_data(fields)
    request = urllib.request.Request(
        f"https://discord.com/api/v10/channels/{thread_id}/messages",
        data=body,
        headers={
            "Authorization": f"Bot {token}",
            "Content-Type": content_type,
            "User-Agent": "legal-redactor/0.1",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=30) as response:
            data = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        exc.read()
        raise DiscordApiError(f"Discord 发送失败: HTTP {exc.code}") from exc
    except OSError as exc:
        raise DiscordApiError("Discord 发送失败: 网络不可达") from exc
    return {
        "message_id": str(data.get("id", "")),
        "channel_id": str(data.get("channel_id", "")),
    }


def _safe_discord_attachment_message(filename: str, message: str = "") -> str:
    safe_filename = Path(filename).name or "redacted.txt"
    default = f"脱敏文件已生成，请见附件：{safe_filename}"
    value = re.sub(r"\r\n?", "\n", str(message)).strip()
    if not value or _contains_local_path_text(value):
        return default[:1900]
    value = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", value)
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return (value or default)[:1900]


def _safe_public_error_message(message: str) -> str:
    return re.sub(r"(?:/Users/|/Volumes/|/private/|~)[^\\s\"'，。；;]+", "<local-path>", message)


def _multipart_form_data(fields: list[tuple[str, str, str | None, bytes]]) -> tuple[bytes, str]:
    boundary = f"----legal-redactor-{next(tempfile._get_candidate_names())}"
    chunks: list[bytes] = []
    for name, content_type, filename, value in fields:
        chunks.append(f"--{boundary}\r\n".encode("ascii"))
        disposition = f'Content-Disposition: form-data; name="{name}"'
        if filename:
            disposition += f'; filename="{filename}"'
        chunks.append(f"{disposition}\r\n".encode("utf-8"))
        chunks.append(f"Content-Type: {content_type}\r\n\r\n".encode("utf-8"))
        chunks.append(value)
        chunks.append(b"\r\n")
    chunks.append(f"--{boundary}--\r\n".encode("ascii"))
    return b"".join(chunks), f"multipart/form-data; boundary={boundary}"


def _suggest_case_location_from_filenames(
    filenames: list[str],
    search_roots: list[Path] | None = None,
    *,
    source_dir: str = "",
    discord_thread_url: str = "",
) -> dict[str, object]:
    return case_suggest_case_location_from_filenames(
        filenames,
        search_roots,
        source_dir=source_dir,
        discord_thread_url=discord_thread_url,
    )


def _best_case_location(
    matches: list[tuple[Path, Path]],
    wanted: set[str],
) -> tuple[Path | None, list[Path]]:
    if not matches:
        return None, []

    scores: dict[Path, set[str]] = {}
    for file_path, case_dir_path in matches:
        scores.setdefault(case_dir_path.resolve(), set()).add(file_path.name)
    ranked = sorted(scores.items(), key=lambda item: (-len(item[1]), str(item[0])))
    if not ranked:
        return None, []

    best_count = len(ranked[0][1])
    best_dirs = [path for path, names in ranked if len(names) == best_count]
    if len(best_dirs) == 1:
        return best_dirs[0], []
    return None, best_dirs


def _case_dir_for_matched_file(root: Path, path: Path) -> Path:
    root = root.resolve()
    path = path.resolve()
    try:
        relative = path.relative_to(root)
    except ValueError:
        return path.parent.resolve()
    if _looks_like_case_root(root) and len(relative.parts) > 1:
        return (root / relative.parts[0]).resolve()
    return path.parent.resolve()


def _looks_like_case_root(root: Path) -> bool:
    try:
        if root.resolve() == default_case_root().resolve():
            return True
    except OSError:
        pass
    return root.name in {"案件资料", "legal-redactor-cases"}


def _case_manifest_fields(case_dir_path: Path) -> dict[str, str]:
    return {
        key: value
        for key, value in manifest_fields_for_case_dir(case_dir_path).items()
        if key in {"discord_thread_url", "discord_thread_id", "workflow_state", "manifest"}
    }


def _case_location_search_roots() -> list[Path]:
    candidates: list[Path] = [
        default_case_root(),
        Path("~/Documents").expanduser(),
        Path("~/Downloads").expanduser(),
        Path("~/Desktop").expanduser(),
    ]
    volumes = Path("/Volumes")
    if volumes.exists():
        for volume in volumes.iterdir():
            if volume.name.startswith("."):
                continue
            candidates.append(volume)
            case_materials = volume / "案件资料"
            if case_materials.exists():
                candidates.insert(0, case_materials)

    seen: set[Path] = set()
    roots: list[Path] = []
    for candidate in candidates:
        try:
            resolved = candidate.expanduser().resolve()
        except OSError:
            continue
        if resolved.exists() and resolved not in seen:
            seen.add(resolved)
            roots.append(resolved)
    return roots


def _find_matching_files(root: Path, wanted: set[str], *, max_depth: int = 5, max_entries: int = 30000) -> list[Path]:
    matches: list[Path] = []
    root = root.resolve()
    visited = 0
    for current, dirs, files in os.walk(root):
        visited += len(dirs) + len(files)
        if visited > max_entries:
            break
        current_path = Path(current)
        depth = len(current_path.relative_to(root).parts)
        dirs[:] = [item for item in dirs if not item.startswith(".") and item not in {"__pycache__", ".git", ".venv"}]
        if depth >= max_depth:
            dirs[:] = []
        for filename in files:
            if filename in wanted:
                matches.append(current_path / filename)
    return matches


def _render_mapping_review_toolbar(redaction_map: RedactionMap, review_candidates: list | None = None) -> str:
    review_texts = _review_candidate_text_set(review_candidates or [])
    counts = {key: 0 for key in MAPPING_REVIEW_CATEGORY_LABELS}
    for entry in redaction_map.mappings:
        for category in _classify_mapping_review_row(entry, review_candidate_texts=review_texts):
            counts[category] += 1
    buttons = [
        f'<button type="button" class="mapping-filter active" data-map-filter="all">'
        f'全部 <span>{len(redaction_map.mappings)}</span></button>'
    ]
    for category, label in MAPPING_REVIEW_CATEGORY_LABELS.items():
        buttons.append(
            f'<button type="button" class="mapping-filter" data-map-filter="{category}">'
            f'{html.escape(label)} <span>{counts[category]}</span></button>'
        )
    return (
        '<div id="mapping-review-toolbar" class="mapping-toolbar">'
        '<div class="mapping-toolbar-head"><strong>复核筛选</strong>'
        '<span class="hint">默认显示全部；点击分类只看需要处理的行。</span></div>'
        f'<div class="mapping-filter-row">{"".join(buttons)}</div>'
        '</div>'
    )


def _render_category_badges(categories: list[str], *, restore_reasons: list[dict[str, str]] | None = None) -> str:
    if not categories:
        return ""
    badges = ""
    for category in categories:
        attrs = ""
        label = MAPPING_REVIEW_CATEGORY_LABELS[category]
        if category == "restore_risk" and restore_reasons:
            codes = ",".join(str(item.get("reason_code") or "") for item in restore_reasons if item.get("reason_code"))
            messages = "；".join(str(item.get("message") or "") for item in restore_reasons if item.get("message"))
            attrs = f' data-restore-risk-codes="{html.escape(codes)}" title="{html.escape(messages)}"'
        badges += (
            f'<span class="row-badge row-badge-{html.escape(category)}"{attrs}>'
            f'{html.escape(label)}</span>'
        )
    return f'<div class="row-tags">{badges}</div>'


def _render_sample_summary_panel() -> str:
    return (
        '<div id="sample-summary-panel" class="sample-summary-panel" hidden>'
        '<strong>样本学习摘要</strong>'
        '<div id="sample-summary-content" class="sample-summary-content"></div>'
        '</div>'
    )


def _review_candidate_texts_json(review_candidates: list | None = None) -> str:
    return json.dumps(sorted(_review_candidate_text_set(review_candidates or [])), ensure_ascii=False)


def _render_mapping_edit_rows(redaction_map: RedactionMap, review_candidates: list | None = None) -> str:
    review_texts = _review_candidate_text_set(review_candidates or [])
    mappings = sort_mapping_entries(list(redaction_map.mappings))
    rows = [
        _render_mapping_edit_row(i, e, review_candidate_texts=review_texts)
        for i, e in enumerate(mappings)
    ]
    rows.append(_render_blank_mapping_row(len(rows)))
    return "".join(rows)


def _render_mapping_edit_row(index: int, entry: MappingEntry, review_candidate_texts: set[str] | None = None) -> str:
    role = entry.role or ""
    reason = entry.reason or ""
    restore = "1" if entry.restore_by_default else "0"
    categories = _classify_mapping_review_row(entry, review_candidate_texts=review_candidate_texts or set())
    category_attr = html.escape(" ".join(categories))
    tags_html = _render_category_badges(categories, restore_reasons=_restore_risk_reasons(entry))
    return f"""
        <tr data-map-row="{index}" data-categories="{category_attr}">
          <td><input name="map_type" value="{html.escape(entry.type)}"></td>
          <td><textarea name="map_original" rows="2">{html.escape(entry.original)}</textarea></td>
          <td><textarea name="map_masked" rows="2">{html.escape(entry.masked)}</textarea></td>
          <td><textarea name="map_reason" rows="2" placeholder="为什么删除/修改/添加">{html.escape(reason)}</textarea></td>
          <td>{html.escape(entry.source)}</td>
          <td>{entry.confidence:.2f}</td>
          <td><label class="inline"><input type="checkbox" name="row_delete" value="{index}"> 删除</label>
            {tags_html}
            <input type="hidden" name="map_role" value="{html.escape(role)}">
            <input type="hidden" name="map_source" value="{html.escape(entry.source)}">
            <input type="hidden" name="map_confidence" value="{entry.confidence}">
            <input type="hidden" name="map_restore_by_default" value="{restore}">
          </td>
        </tr>
    """


def _render_blank_mapping_row(index: int) -> str:
    return f"""
        <tr data-map-row="{index}" data-categories="">
          <td><input name="map_type" value="manual" placeholder="person/org"></td>
          <td><textarea name="map_original" rows="2" placeholder="新增要替换的原文"></textarea></td>
          <td><textarea name="map_masked" rows="2" placeholder="替换为"></textarea></td>
          <td><textarea name="map_reason" rows="2" placeholder="为什么新增这条"></textarea></td>
          <td>manual</td>
          <td>1.0</td>
          <td><label class="inline"><input type="checkbox" name="row_delete" value="{index}"> 删除</label>
            <input type="hidden" name="map_role" value="">
            <input type="hidden" name="map_source" value="manual">
            <input type="hidden" name="map_confidence" value="1.0">
            <input type="hidden" name="map_restore_by_default" value="1">
          </td>
        </tr>
    """


async def _read_input_documents(
    text: str,
    file: UploadFile | None,
    files: list[UploadFile],
    case_folder_files: list[UploadFile] | None = None,
) -> list[InputDocument]:
    documents = []
    if text.strip():
        documents.append(InputDocument(source_file="粘贴文本.txt", text=text))

    target_files = []
    if file and file.filename: target_files.append(file)
    if files: target_files.extend([f for f in files if f.filename])
    folder_target_files = [
        item
        for item in (case_folder_files or [])
        if item.filename and _is_supported_folder_upload_filename(item.filename)
    ]

    for item in target_files:
        try:
            content = await _read_upload_text(item)
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"读取文件 {item.filename} 失败: {exc}") from exc
        documents.append(InputDocument(source_file=item.filename, text=content))

    for item in folder_target_files:
        try:
            content = await _read_upload_text(item)
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"读取文件 {item.filename} 失败: {exc}") from exc
        documents.append(InputDocument(source_file=item.filename, text=content))

    if not documents: raise ValueError("未提供任何待脱敏的文本或文件")
    return documents

def _decode_text_bytes(data: bytes, filename: str) -> str:
    """尝试以不同编码解析上传的二进制文本字节流，主要支持 UTF-8, GB18030, GBK 等。"""
    for encoding in ("utf-8-sig", "gb18030", "gbk", "utf-8", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _suffix_for_filename(filename: str) -> str:
    return "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ".txt"


def _is_supported_folder_upload_filename(filename: str) -> bool:
    name = PurePosixPath(str(filename).replace("\\", "/")).name
    return bool(name and not name.startswith("._") and _suffix_for_filename(name) in SUPPORTED_UPLOAD_SUFFIXES)


def _docx_bytes_to_text(data: bytes) -> str:
    from docx import Document
    doc = Document(BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    return "\n".join(texts)


def _legacy_doc_bytes_to_text(data: bytes, filename: str) -> str:
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        try:
            result = subprocess.run(
                [
                    "/usr/bin/textutil",
                    "-convert",
                    "txt",
                    "-stdout",
                    "-encoding",
                    "UTF-8",
                    "--",
                    str(tmp_path),
                ],
                check=False,
                capture_output=True,
                timeout=60,
            )
        except FileNotFoundError as exc:
            raise ValueError("读取 .doc 需要 macOS textutil，请先用 Word/WPS 另存为 .docx 或导出为 .txt") from exc
        except subprocess.TimeoutExpired as exc:
            raise ValueError(f"读取文件 {filename} 失败: .doc 转文本超时，请先另存为 .docx 或 .txt") from exc
        if result.returncode != 0:
            error = result.stderr.decode("utf-8", errors="replace").strip()
            detail = f": {error}" if error else ""
            raise ValueError(f"读取文件 {filename} 失败: .doc 转文本失败{detail}。请先另存为标准 .docx 或 .txt")
        return result.stdout.decode("utf-8", errors="replace")
    finally:
        try:
            tmp_path.unlink()
        except OSError:
            pass


async def _read_restore_map_text(map_json: str, map_file: UploadFile | None) -> str:
    """读取还原映射表，支持直接粘贴 JSON 和上传 JSON 文件（包括加密映射表）。"""
    map_text = ""
    if map_file and map_file.filename:
        data = await map_file.read()
        try:
            map_text = _decode_text_bytes(data, map_file.filename)
            json.loads(map_text)
        except (json.JSONDecodeError, UnicodeDecodeError):
            from ._crypto import decrypt
            map_text = decrypt(data)
    elif map_json.strip():
        map_text = map_json

    if not map_text:
        raise ValueError("请提供有效的映射表内容或文件")
    return map_text

async def _read_upload_text(file: UploadFile) -> str:
    data = await file.read()
    suffix = _suffix_for_filename(file.filename)
    if suffix in (".txt", ".md"):
        return _decode_text_bytes(data, file.filename)
    if suffix == ".docx":
        try:
            return _docx_bytes_to_text(data)
        except BadZipFile as exc:
            raise ValueError(
                f"读取文件 {file.filename} 失败: 该文件不是有效的 .docx。"
                "如果它是旧版 .doc、RTF、WPS 格式或文件已损坏，请先用 Word/WPS 另存为标准 .docx，或导出为 .txt 后再上传。"
            ) from exc
        except Exception as exc:
            raise ValueError(f"读取文件 {file.filename} 失败: {exc}") from exc
    if suffix == ".doc":
        return _legacy_doc_bytes_to_text(data, file.filename)
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ValueError("读取 pdf 需要安装 pypdf：pip install pypdf") from exc
        try:
            reader = PdfReader(BytesIO(data))
            text = []
            for page in reader.pages:
                text.append(page.extract_text() or "")
            return "\n".join(text)
        except Exception as exc:
            raise ValueError(f"读取文件 {file.filename} 失败: PDF 格式无效或文件已损坏") from exc
    return ""

def _redaction_map_from_rows(
    version: str, created_at: str, mode: str, source_file: str,
    map_type: list[str], map_original: list[str], map_masked: list[str],
    map_role: list[str], map_source: list[str], map_confidence: list[str],
    map_reason: list[str], map_restore_by_default: list[str], row_delete: list[str],
) -> RedactionMap:
    deleted = set(row_delete)
    row_count = max(len(map_original), len(map_masked), len(map_type))
    mappings: list[MappingEntry] = []
    for index in range(row_count):
        if str(index) in deleted: continue
        original = _form_list_value(map_original, index).strip()
        masked = _form_list_value(map_masked, index).strip()
        if not original or not masked: continue
        role = _form_list_value(map_role, index).strip() or None
        try:
            confidence = float(_form_list_value(map_confidence, index) or "1.0")
        except ValueError:
            confidence = 1.0
        mappings.append(MappingEntry(
            type=_form_list_value(map_type, index).strip() or "manual",
            original=original, masked=masked, role=role,
            source=_form_list_value(map_source, index).strip() or "manual",
            confidence=confidence,
            restore_by_default=_form_list_value(map_restore_by_default, index) != "0",
            reason=_form_list_value(map_reason, index).strip() or None,
        ))
    return RedactionMap(
        version=version or "1.0",
        created_at=created_at,
        mode=mode or "normal",
        source_file=source_file or None,
        mappings=sort_mapping_entries(mappings),
    )


def _find_mapping_by_original(mappings: list[MappingEntry], original: str) -> MappingEntry | None:
    value = original.strip()
    for entry in mappings:
        if entry.original == value:
            return entry
    return None


_ORG_ALIAS_SUFFIXES = (
    "有限责任公司",
    "股份有限公司",
    "集团有限公司",
    "有限公司",
    "公司",
    "集团",
    "银行",
    "信用社",
    "合作社",
    "事务所",
    "律所",
    "学校",
    "医院",
    "法院",
    "检察院",
    "委员会",
    "村委会",
    "居委会",
    "商行",
    "经营部",
    "店",
    "厂",
    "机构",
)


def _organization_originals_are_aliases(left: str, right: str) -> bool:
    a = left.strip()
    b = right.strip()
    if not a or not b:
        return False
    if a == b:
        return True
    shorter, longer = (a, b) if len(a) <= len(b) else (b, a)
    if longer.startswith(shorter):
        tail = longer[len(shorter):]
        if not tail or tail in _ORG_ALIAS_SUFFIXES:
            return True
    for full_name in (a, b):
        if full_name.endswith(_ORG_ALIAS_SUFFIXES):
            cores = derived_organization_alias_cores(full_name)
            other = b if full_name == a else a
            if other in cores or f"{other}公司" == full_name or other == full_name.replace("公司", ""):
                return True
    return False


def _mapping_entries_share_entity(left: MappingEntry, right: MappingEntry) -> bool:
    if left.type == right.type == "person":
        return left.original.strip() == right.original.strip()
    if left.type in {"organization", "individual_business"} and right.type in {"organization", "individual_business"}:
        return _organization_originals_are_aliases(left.original, right.original)
    return left.type == right.type and left.original.strip() == right.original.strip()


def _mapping_entity_group_ids(mappings: list[MappingEntry]) -> list[int]:
    parent = list(range(len(mappings)))

    def find(index: int) -> int:
        while parent[index] != index:
            parent[index] = parent[parent[index]]
            index = parent[index]
        return index

    def union(left: int, right: int) -> None:
        root_left, root_right = find(left), find(right)
        if root_left != root_right:
            parent[root_right] = root_left

    for left in range(len(mappings)):
        for right in range(left + 1, len(mappings)):
            if _mapping_entries_share_entity(mappings[left], mappings[right]):
                union(left, right)

    leaders: dict[int, int] = {}
    group_ids: list[int] = []
    for index in range(len(mappings)):
        root = find(index)
        if root not in leaders:
            leaders[root] = len(leaders)
        group_ids.append(leaders[root])
    return group_ids


def _renumber_mapping_placeholders(mappings: list[MappingEntry]) -> list[MappingEntry]:
    if not mappings:
        return []
    group_ids = _mapping_entity_group_ids(mappings)
    members: dict[int, list[int]] = {}
    for index, group_id in enumerate(group_ids):
        members.setdefault(group_id, []).append(index)

    ordered_group_ids = sorted(members, key=lambda group_id: min(members[group_id]))
    group_ordinals: dict[int, str] = {}
    type_counts: dict[str, int] = {}
    person_counts: dict[str, int] = {}
    for group_id in ordered_group_ids:
        representative = mappings[members[group_id][0]]
        group_ordinals[group_id] = _next_group_ordinal(representative, type_counts, person_counts)

    renumbered: list[MappingEntry] = []
    for index, entry in enumerate(mappings):
        masked = _mask_with_group_ordinal(entry, group_ordinals[group_ids[index]])
        renumbered.append(replace(entry, masked=masked) if masked != entry.masked else entry)
    return renumbered


def _next_group_ordinal(
    entry: MappingEntry,
    type_counts: dict[str, int],
    person_counts: dict[str, int],
) -> str:
    if entry.type == "person":
        stem = _person_mask_stem(entry)
        person_counts[stem] = person_counts.get(stem, 0) + 1
        return _ordinal_value(person_counts[stem])
    counter_key = _renumber_counter_key(entry)
    type_counts[counter_key] = type_counts.get(counter_key, 0) + 1
    return _ordinal_value(type_counts[counter_key])


def _renumber_counter_key(entry: MappingEntry) -> str:
    if entry.type in {"organization", "individual_business"}:
        return "organization"
    if entry.type in {"location", "grassroots_org"}:
        return "location"
    if entry.type == "project":
        return "project"
    return entry.type or "manual"


def _person_mask_stem(entry: MappingEntry) -> str:
    match = re.match(r"^(.+?某)(?:[甲乙丙丁戊己庚辛壬癸]|\d+)$", entry.masked or "")
    if match:
        return match.group(1)
    if re.fullmatch(r"[\u4e00-\u9fa5]{2,4}", entry.original or ""):
        return f"{entry.original[0]}某"
    return "自然人"


def _mask_with_group_ordinal(entry: MappingEntry, ordinal: str) -> str:
    if entry.type == "person":
        return f"{_person_mask_stem(entry)}{ordinal}"
    if entry.type in {"organization", "individual_business"}:
        return _mask_with_ordinal_prefix(entry, ordinal, _manual_organization_suffix)
    if entry.type in {"location", "grassroots_org"}:
        return _mask_with_ordinal_prefix(entry, ordinal, _manual_location_suffix)
    if entry.type == "project":
        return _mask_with_ordinal_prefix(entry, ordinal, _project_suffix)
    return entry.masked


def _mask_with_ordinal_prefix(entry: MappingEntry, ordinal: str, suffix_from_original) -> str:
    match = re.match(r"^([甲乙丙丁戊己庚辛壬癸]|\d+)(.*)$", entry.masked or "")
    if match:
        suffix = match.group(2)
        if not suffix:
            return ordinal
        return f"{ordinal}{suffix}"
    return f"{ordinal}{suffix_from_original(entry.original)}"


def _project_suffix(original: str) -> str:
    if original.endswith("工程"):
        return "工程"
    return "项目"


def _suggest_manual_mapping_entry(original: str, entity_type: str, existing: list[MappingEntry]) -> MappingEntry:
    value = original.strip()
    masked = _suggest_manual_mask(value, entity_type, existing)
    return MappingEntry(
        type=entity_type,
        original=value,
        masked=masked,
        role=None,
        source="manual_selection",
        confidence=1.0,
        restore_by_default=True,
    )


def _suggest_manual_mask(original: str, entity_type: str, existing: list[MappingEntry]) -> str:
    if entity_type == "person":
        if re.fullmatch(r"[\u4e00-\u9fa5]{2,4}", original):
            surname = original[0]
            return f"{surname}某{_next_person_ordinal(surname, existing)}"
        return f"自然人{_next_mask_ordinal(existing, {'person'}, '自然人')}"
    if entity_type == "organization":
        suffix = _manual_organization_suffix(original)
        return f"{_next_mask_ordinal(existing, {'organization', 'individual_business'}, '')}{suffix}"
    if entity_type == "location":
        suffix = _manual_location_suffix(original)
        return f"{_next_mask_ordinal(existing, {'location', 'grassroots_org'}, '')}{suffix}"
    return f"敏感信息{_next_mask_ordinal(existing, {entity_type}, '敏感信息')}"


def _next_person_ordinal(surname: str, existing: list[MappingEntry]) -> str:
    used = 0
    pattern = re.compile(rf"^{re.escape(surname)}某([甲乙丙丁戊己庚辛壬癸]|\d+)$")
    for entry in existing:
        if entry.type != "person":
            continue
        match = pattern.match(entry.masked)
        if not match:
            continue
        used = max(used, _ordinal_index(match.group(1)))
    return _ordinal_value(used + 1)


def _next_mask_ordinal(existing: list[MappingEntry], entity_types: set[str], prefix: str) -> str:
    used = 0
    pattern = re.compile(rf"^{re.escape(prefix)}([甲乙丙丁戊己庚辛壬癸]|\d+)")
    for entry in existing:
        if entry.type not in entity_types:
            continue
        match = pattern.match(entry.masked)
        if match:
            used = max(used, _ordinal_index(match.group(1)))
    return _ordinal_value(used + 1)


def _ordinal_index(value: str) -> int:
    if value in CN_ORDINALS:
        return CN_ORDINALS.index(value) + 1
    try:
        return int(value)
    except ValueError:
        return 0


def _ordinal_value(index: int) -> str:
    if index <= len(CN_ORDINALS):
        return CN_ORDINALS[index - 1]
    return str(index)


def _manual_organization_suffix(original: str) -> str:
    for suffix in (
        "有限公司", "股份有限公司", "公司", "集团", "银行", "信用社", "合作社",
        "事务所", "律所", "学校", "医院", "法院", "检察院", "委员会",
        "村委会", "居委会", "商行", "经营部", "店", "厂",
    ):
        if original.endswith(suffix):
            if suffix in {"有限公司", "股份有限公司"}:
                return "公司"
            return suffix
    return "机构"


def _manual_location_suffix(original: str) -> str:
    for suffix in (
        "自治区", "自治州", "居民委员会", "村民委员会", "街道", "社区",
        "省", "市", "区", "县", "旗", "镇", "乡", "村", "小区", "项目", "工程",
    ):
        if original.endswith(suffix):
            if suffix in {"居民委员会", "村民委员会"}:
                return suffix
            return suffix
    return "地"


def _form_list_value(values: list[str], index: int) -> str:
    if index >= len(values): return ""
    return values[index]


def _data_download(filename: str, mime: str, content: str) -> str:
    return f"data:{mime};charset=utf-8,{urllib.parse.quote(content)}"


def _binary_download(mime: str, content: bytes) -> str:
    return f"data:{mime};base64,{base64.b64encode(content).decode('ascii')}"


def _documents_bundle_json(documents: list[RedactedDocument]) -> str:
    return json.dumps([{"source_file": d.source_file, "text": d.original_text} for d in documents], ensure_ascii=False)


def _documents_from_bundle_json(value: str) -> list[InputDocument]:
    if not value.strip(): return []
    try: payload = json.loads(value)
    except json.JSONDecodeError: return []
    if not isinstance(payload, list): return []
    return [InputDocument(source_file=str(i.get("source_file","")), text=str(i.get("text",""))) for i in payload if isinstance(i, dict)]


def _apply_map_to_documents(pipeline: RedactionPipeline, documents: list[InputDocument], redaction_map: RedactionMap) -> list[RedactedDocument]:
    return [RedactedDocument(source_file=d.source_file, original_text=d.text, redacted_text=pipeline.apply_redaction_map(d.text, redaction_map), leaks=pipeline.scan_high_risk_leaks(pipeline.apply_redaction_map(d.text, redaction_map))) for d in documents]


def _highlight_replaced_text(text: str, entries: list[MappingEntry], *, reverse: bool = False) -> str:
    """生成带 <mark> 高亮的 HTML 文本。

    reverse=False：高亮原文中被替换的词（title 显示替换后的内容）。
    reverse=True：高亮脱敏文本中的占位符（title 显示原文）。
    """
    if not entries:
        return html.escape(text)

    spans: list[tuple[int, int, str, str]] = []
    occupied: list[tuple[int, int]] = []

    sorted_entries = sorted(
        entries, key=lambda e: len(e.masked if reverse else e.original), reverse=True
    )

    for entry in sorted_entries:
        search_key = entry.masked if reverse else entry.original
        if not search_key:
            continue
        tooltip = entry.original if reverse else entry.masked

        pos = 0
        while True:
            idx = text.find(search_key, pos)
            if idx < 0:
                break
            end = idx + len(search_key)
            if not any(not (end <= occ_start or idx >= occ_end) for occ_start, occ_end in occupied):
                spans.append((idx, end, search_key, tooltip))
                occupied.append((idx, end))
            pos = idx + 1

    if not spans:
        return html.escape(text)

    spans.sort(key=lambda s: s[0])

    parts: list[str] = []
    last = 0
    for start, end, display, tooltip_text in spans:
        if start > last:
            parts.append(html.escape(text[last:start]))
        title = f"原文：{tooltip_text}" if reverse else f"→ {tooltip_text}"
        parts.append(f'<mark title="{html.escape(title)}">{html.escape(display)}</mark>')
        last = end

    if last < len(text):
        parts.append(html.escape(text[last:]))

    return "".join(parts)


def _guess_location_mask(text: str) -> str:
    """为无后缀的纯地名（如'石家庄''沧州'）猜测合适的掩码。"""
    # 如果原本就是带后缀的，直接返回通用掩码
    for sfx, mask in [("自治区", "某自治区"), ("自治州", "某自治州"),
                       ("街道", "某街道"), ("省", "某省"), ("市", "某市"),
                       ("区", "某区"), ("县", "某县"), ("镇", "某镇"),
                       ("乡", "某乡"), ("村", "某村")]:
        if text.endswith(sfx):
            return mask
    # 无后缀的地名简称：根据常见模式推测
    if re.fullmatch(r"[一-龥]{2,4}", text):
        return "某市"  # 最常见的地名简称是城市名
    return f"地点"


def _simple_mask(text: str, counters: TypeCounters) -> str:
    """为用户确认的实体生成简单掩码，供增量脱敏使用。"""
    # 公司/机构
    if re.search(r"(公司|集团|厂|店|经营部|商行|事务所|律所|银行|信用社)$", text):
        return f"公司{counters.next('company')}"
    # 地名后缀（必须在自然人检查之前，避免"河南省"被误判为姓名）
    # 注意：长后缀必须在短后缀之前检查，避免"内蒙古自治区"被"区"截断
    if text.endswith("自治区"):
        return "某自治区"
    if text.endswith("自治州"):
        return "某自治州"
    if text.endswith("街道"):
        return "某街道"
    if text.endswith("省"):
        return "某省"
    if text.endswith("市"):
        return "某市"
    if text.endswith("区"):
        return "某区"
    if text.endswith("县"):
        return "某县"
    if text.endswith("镇"):
        return "某镇"
    if text.endswith("乡"):
        return "某乡"
    if text.endswith("村"):
        return "某村"
    # 法院
    if "法院" in text:
        return "某法院"
    # 自然人（2-4字，不含地名后缀）
    if re.fullmatch(r"[一-龥]{2,4}", text):
        return f"自然人{counters.next('person')}"
    # 项目/工程
    if "工程" in text or "项目" in text:
        return f"项目{counters.next('project')}"
    # 其他
    return f"敏感信息{counters.next('other')}"
