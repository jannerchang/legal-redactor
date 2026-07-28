from __future__ import annotations

import asyncio
import http.client
import html
import json
import os
import re
import secrets
import subprocess
import urllib.error
import urllib.parse
import urllib.request
import base64
import tempfile
from dataclasses import dataclass, replace
from datetime import datetime
from io import BytesIO
from pathlib import Path, PurePosixPath
from typing import Any
from zipfile import BadZipFile

from . import deps
from .deps import (
    CN_ORDINALS,
    DEFAULT_MODEL_ID,
    EXCEL_INPUT_SUFFIXES,
    ExcelFormulaLeakError,
    File,
    Form,
    HTMLResponse,
    JSONResponse,
    MappingEntry,
    PipelineConfig,
    RecognitionRunStats,
    RedactedDocument,
    RedactionMap,
    Request,
    TypeCounters,
    UploadFile,
    XLSX_MEDIA_TYPE,
    _filter_noise_entity_mappings,
    _page,
    derived_organization_alias_cores,
    extract_workbook_text,
    is_noise_entity_text,
    preview_restore,
    redact_workbook,
    redaction_map_from_json,
    redaction_map_to_json,
    render_batch_redaction_result_page,
    render_home_page,
    render_redaction_result_page,
    render_status_panel,
    restore_docx,
    sort_mapping_entries,
    MAPPING_REVIEW_CATEGORY_LABELS,
    RESTORE_RISK_REASON_LABELS,
)

from ..cases import (
    CaseError,
    InvalidDiscordThreadError,
    InvalidWorkflowInputError,
    case_dir,
    create_or_update_manifest,
    case_root_from_source_dir,
    case_workflow_public,
    case_workflow_state,
    default_case_root,
    load_manifest,
    manifest_fields_for_case_dir,
    parse_discord_thread_id,
    persist_case_redaction,
    raise_for_forged_workflow_fields,
    record_hermes_thread_request,
    suggest_case_location_from_filenames,
    case_location_search_roots,
    case_thread_binding_status,
    validate_case_folder_name,
    workflow_state_message,
)


from .models import InputDocument, SUPPORTED_UPLOAD_SUFFIXES


def _excel_source(document: InputDocument) -> bool:
    return document.source_bytes is not None and document.source_suffix in EXCEL_INPUT_SUFFIXES



def _render_output_document(
    source: InputDocument,
    redacted: RedactedDocument,
    redaction_map: RedactionMap,
    pipeline: deps.RedactionPipeline,
) -> RedactedDocument:
    if not _excel_source(source):
        return redacted
    output_bytes = redact_workbook(source.source_bytes or b"", source.source_file, redaction_map, pipeline.apply_mappings)
    output_filename = f"{PurePosixPath(source.source_file.replace(chr(92), chr(47))).stem or 'redacted'}.redacted.xlsx"
    output_text = extract_workbook_text(output_bytes, output_filename)
    return replace(
        redacted,
        redacted_text=output_text,
        output_filename=output_filename,
        output_media_type=XLSX_MEDIA_TYPE,
        output_bytes=output_bytes,
        leaks=pipeline.scan_high_risk_leaks(output_text),
    )



def _excel_warnings(documents: list[InputDocument]) -> list[str]:
    if any(d.source_suffix == ".xlsm" for d in documents):
        return ["已移除 XLSM 宏并输出为 XLSX；宏、批注、超链接目标和图形文本未参与脱敏。"]
    if any(_excel_source(d) for d in documents):
        return ["Excel 仅脱敏单元格文本；批注、超链接目标和图形文本未参与脱敏。"]
    return []



async def _read_input_documents(
    text: str,
    file: UploadFile | None,
    files: list[UploadFile],
    case_folder_files: list[UploadFile] | None = None,
) -> list[InputDocument]:
    documents: list[InputDocument] = []
    if text.strip():
        documents.append(InputDocument(source_file="粘贴文本.txt", text=text))
    target_files: list[UploadFile] = []
    if file and file.filename:
        target_files.append(file)
    target_files.extend(f for f in files if f.filename)
    target_files.extend(
        item for item in (case_folder_files or [])
        if item.filename and _is_supported_folder_upload_filename(item.filename)
    )
    for item in target_files:
        name = str(item.filename)
        data = await item.read()
        suffix = _suffix_for_filename(name)
        try:
            if suffix in EXCEL_INPUT_SUFFIXES:
                content = extract_workbook_text(data, name)
                documents.append(InputDocument(name, content, data, suffix))
            else:
                documents.append(InputDocument(name, _read_upload_text_from_bytes(data, name)))
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"读取文件 {name} 失败: {exc}") from exc
    if not documents:
        raise ValueError("未提供任何待脱敏的文本或文件")
    return documents



def _decode_text_bytes(data: bytes, filename: str) -> str:
    """尝试以不同编码解析上传的二进制文本字节流，主要支持 UTF-8, GB18030, GBK 等。"""
    for encoding in ("utf-8-sig", "gb18030", "gbk", "utf-8", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")



def _suffix_for_filename(filename: str) -> str:
    return "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ".txt"



def _is_supported_folder_upload_filename(filename: str) -> bool:
    name = PurePosixPath(str(filename).replace("\\", "/")).name
    return bool(name and not name.startswith("._") and _suffix_for_filename(name) in SUPPORTED_UPLOAD_SUFFIXES)



def _docx_bytes_to_text(data: bytes) -> str:
    from docx import Document
    doc = Document(BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    return "\n".join(texts)



def _legacy_doc_bytes_to_text(data: bytes, filename: str) -> str:
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        try:
            result = subprocess.run(
                [
                    "/usr/bin/textutil",
                    "-convert",
                    "txt",
                    "-stdout",
                    "-encoding",
                    "UTF-8",
                    "--",
                    str(tmp_path),
                ],
                check=False,
                capture_output=True,
                timeout=60,
            )
        except FileNotFoundError as exc:
            raise ValueError("读取 .doc 需要 macOS textutil，请先用 Word/WPS 另存为 .docx 或导出为 .txt") from exc
        except subprocess.TimeoutExpired as exc:
            raise ValueError(f"读取文件 {filename} 失败: .doc 转文本超时，请先另存为 .docx 或 .txt") from exc
        if result.returncode != 0:
            error = result.stderr.decode("utf-8", errors="replace").strip()
            detail = f": {error}" if error else ""
            raise ValueError(f"读取文件 {filename} 失败: .doc 转文本失败{detail}。请先另存为标准 .docx 或 .txt")
        return result.stdout.decode("utf-8", errors="replace")
    finally:
        try:
            tmp_path.unlink()
        except OSError:
            pass



async def _read_restore_map_text(map_json: str, map_file: UploadFile | None) -> str:
    """读取还原映射表，支持直接粘贴 JSON 和上传 JSON 文件（包括加密映射表）。"""
    map_text = ""
    if map_file and map_file.filename:
        data = await map_file.read()
        try:
            map_text = _decode_text_bytes(data, map_file.filename)
            json.loads(map_text)
        except (json.JSONDecodeError, UnicodeDecodeError):
            from .._crypto import decrypt
            map_text = decrypt(data)
    elif map_json.strip():
        map_text = map_json

    if not map_text:
        raise ValueError("请提供有效的映射表内容或文件")
    return map_text



def _read_upload_text_from_bytes(data: bytes, filename: str) -> str:
    suffix = _suffix_for_filename(filename)
    if suffix in (".txt", ".md"):
        return _decode_text_bytes(data, filename)
    if suffix == ".docx":
        try:
            return _docx_bytes_to_text(data)
        except BadZipFile as exc:
            raise ValueError(
                f"读取文件 {filename} 失败: 该文件不是有效的 .docx。"
                "如果它是旧版 .doc、RTF、WPS 格式或文件已损坏，请先用 Word/WPS 另存为标准 .docx，或导出为 .txt 后再上传。"
            ) from exc
        except Exception as exc:
            raise ValueError(f"读取文件 {filename} 失败: {exc}") from exc
    if suffix == ".doc":
        return _legacy_doc_bytes_to_text(data, filename)
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ValueError("读取 pdf 需要安装 pypdf：pip install pypdf") from exc
        try:
            reader = PdfReader(BytesIO(data))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise ValueError(f"读取文件 {filename} 失败: PDF 格式无效或文件已损坏") from exc
    return ""



async def _read_upload_text(file: UploadFile) -> str:
    data = await file.read()
    return _read_upload_text_from_bytes(data, file.filename)



def _form_list_value(values: list[str], index: int) -> str:
    if index >= len(values):
        return ""
    return values[index]



def _data_download(filename: str, mime: str, content: str) -> str:
    prefix = "\ufeff" if mime == "text/plain" and filename.lower().endswith(".txt") else ""
    return f"data:{mime};charset=utf-8,{urllib.parse.quote(prefix + content)}"



def _binary_download(mime: str, content: bytes) -> str:
    return f"data:{mime};base64,{base64.b64encode(content).decode('ascii')}"



def _documents_bundle_json(documents: list[RedactedDocument], sources: list[InputDocument] | None = None) -> str:
    source_items = sources or [InputDocument(d.source_file, d.original_text) for d in documents]
    payload = []
    for source in source_items:
        source_suffix = source.source_suffix if source.source_suffix in EXCEL_INPUT_SUFFIXES else ".txt"
        payload.append({
            "source_file": source.source_file,
            "text": source.text,
            "source_suffix": source_suffix,
            "source_base64": base64.b64encode(source.source_bytes).decode("ascii") if source.source_bytes is not None else "",
        })
    return json.dumps(payload, ensure_ascii=False)



def _documents_from_bundle_json(value: str) -> list[InputDocument]:
    if not value.strip():
        return []
    try:
        payload = json.loads(value)
    except json.JSONDecodeError as exc:
        raise ValueError("Excel 源文件状态无效，请重新上传") from exc
    if not isinstance(payload, list):
        raise ValueError("Excel 源文件状态无效，请重新上传")
    documents: list[InputDocument] = []
    for item in payload:
        if not isinstance(item, dict):
            raise ValueError("Excel 源文件状态无效，请重新上传")
        source_file = str(item.get("source_file") or "")
        text = str(item.get("text") or "")
        suffix = str(item.get("source_suffix") or ".txt").lower()
        encoded = str(item.get("source_base64") or "")
        if suffix in EXCEL_INPUT_SUFFIXES:
            if not encoded or not source_file.lower().endswith(suffix):
                raise ValueError("Excel 源文件状态无效，请重新上传")
            try:
                source_bytes = base64.b64decode(encoded, validate=True)
                if extract_workbook_text(source_bytes, source_file) != text:
                    raise ValueError
            except Exception as exc:
                raise ValueError("Excel 源文件状态无效，请重新上传") from exc
            documents.append(InputDocument(source_file, text, source_bytes, suffix))
        else:
            documents.append(InputDocument(source_file, text))
    return documents



def _apply_map_to_documents(pipeline: deps.RedactionPipeline, documents: list[InputDocument], redaction_map: RedactionMap) -> list[RedactedDocument]:
    result: list[RedactedDocument] = []
    for document in documents:
        redacted_text = pipeline.apply_redaction_map(document.text, redaction_map)
        result.append(RedactedDocument(
            source_file=document.source_file,
            original_text=document.text,
            redacted_text=redacted_text,
            leaks=pipeline.scan_high_risk_leaks(redacted_text),
        ))
    return result
