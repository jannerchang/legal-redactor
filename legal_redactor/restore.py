from __future__ import annotations

import difflib
from pathlib import Path

from ._replace_text import replace_by_lookup
from .models import MappingEntry, RedactionMap, RestorePreview


def restore_text(redacted_text: str, redaction_map: RedactionMap, restore_all: bool = True) -> str:
    """还原脱敏文本：按映射表将占位符替换回原文。

    使用位置查找 + 字符数组倒序替换，避免全局 str.replace()
    可能的误匹配问题（例如某个占位符碰巧出现在不应还原的上下文中）。
    """
    entries = _entries_to_restore(redaction_map.mappings, restore_all)
    lookup = _restore_lookup(entries)
    return replace_by_lookup(redacted_text, lookup)


def restore_docx(
    input_path: str | Path,
    output_path: str | Path,
    redaction_map: RedactionMap,
    restore_all: bool = True,
) -> int:
    """按映射表还原 docx，并尽量保留原 Word 结构和格式。

    替换会遍历正文、表格、页眉和页脚。若占位符被 Word 拆成多个 run，
    还原文字会使用占位符起始 run 的格式，其他未命中的文字格式保持不变。
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("还原 docx 需要安装 python-docx：pip install -r requirements.txt") from exc

    entries = _entries_to_restore(redaction_map.mappings, restore_all)
    lookup = _restore_lookup(entries)
    document = Document(str(input_path))
    count = 0
    for paragraph in _iter_document_paragraphs(document):
        count += _restore_paragraph_runs(paragraph, lookup)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return count


def preview_restore(redacted_text: str, redaction_map: RedactionMap, restore_all: bool = True) -> RestorePreview:
    restored_entries = _entries_to_restore(redaction_map.mappings, restore_all)
    skipped_entries = [entry for entry in redaction_map.mappings if entry not in restored_entries]
    restored_text = restore_text(redacted_text, redaction_map, restore_all)
    diff = "\n".join(
        difflib.unified_diff(
            redacted_text.splitlines(),
            restored_text.splitlines(),
            fromfile="redacted",
            tofile="restored",
            lineterm="",
        )
    )
    return RestorePreview(
        restored_text=restored_text,
        restored_entries=restored_entries,
        skipped_entries=skipped_entries,
        diff=diff,
    )


def _restore_lookup(entries: list[MappingEntry]) -> dict[str, str]:
    """Restore unique masks exactly and shared registry masks to their primary text."""
    grouped: dict[str, list[MappingEntry]] = {}
    for entry in entries:
        if entry.masked:
            grouped.setdefault(entry.masked, []).append(entry)

    lookup: dict[str, str] = {}
    for masked, candidates in grouped.items():
        if len(candidates) == 1:
            lookup[masked] = candidates[0].original
            continue
        canonical = next(
            (
                entry.restore_original
                for entry in candidates
                if entry.restore_original
            ),
            None,
        )
        lookup[masked] = canonical if canonical is not None else candidates[-1].original
    return lookup




def _entries_to_restore(entries: list[MappingEntry], restore_all: bool) -> list[MappingEntry]:
    # 保留参数以兼容旧调用；统一还原标准始终恢复映射表中的全部条目。
    _ = restore_all
    return list(entries)


def _iter_document_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        yield from section.header.paragraphs
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)
        yield from section.footer.paragraphs
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _restore_paragraph_runs(paragraph, lookup: dict[str, str]) -> int:
    if not lookup or not paragraph.runs:
        return 0

    text = paragraph.text
    if not text:
        return 0

    matches = []
    for masked, original in lookup.items():
        start = text.find(masked)
        while start != -1:
            matches.append((start, start + len(masked), original))
            start = text.find(masked, start + len(masked))
    if not matches:
        return 0

    matches.sort(key=lambda item: (item[0], -(item[1] - item[0])))
    non_overlapping = []
    last_end = -1
    for match in matches:
        if match[0] >= last_end:
            non_overlapping.append(match)
            last_end = match[1]

    run_chars = [list(run.text) for run in paragraph.runs]
    char_positions: list[tuple[int, int]] = []
    for run_index, chars in enumerate(run_chars):
        for char_index, _ in enumerate(chars):
            char_positions.append((run_index, char_index))

    for start, end, original in reversed(non_overlapping):
        covered = char_positions[start:end]
        if not covered:
            continue
        first_run, _ = covered[0]
        covered_by_run: dict[int, list[int]] = {}
        for run_index, char_index in covered:
            covered_by_run.setdefault(run_index, []).append(char_index)

        first_indices = covered_by_run.pop(first_run, [])
        if first_indices:
            first_indices.sort()
            first = first_indices[0]
            last = first_indices[-1]
            run_chars[first_run][first : last + 1] = list(original)

        for run_index, indices in covered_by_run.items():
            for char_index in sorted(indices, reverse=True):
                del run_chars[run_index][char_index]

    for run, chars in zip(paragraph.runs, run_chars):
        run.text = "".join(chars)
    return len(non_overlapping)
