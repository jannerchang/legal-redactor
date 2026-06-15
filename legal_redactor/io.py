from __future__ import annotations

import json
import os
from pathlib import Path

from .models import RedactionMap


SUPPORTED_INPUT_SUFFIXES = {".txt", ".md", ".docx", ".pdf"}


def read_document(path: str | Path) -> str:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8")
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("读取 docx 需要安装 python-docx：pip install -r requirements.txt") from exc
        document = Document(str(file_path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("读取 pdf 需要安装 pypdf：pip install pypdf") from exc
        reader = PdfReader(str(file_path))
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or "")
        return "\n".join(text)
    raise ValueError(f"不支持的输入格式：{suffix}")


def write_document(path: str | Path, text: str) -> None:
    file_path = Path(path)
    file_path.parent.mkdir(parents=True, exist_ok=True)
    suffix = file_path.suffix.lower()
    if suffix in {".txt", ".md"}:
        file_path.write_text(text, encoding="utf-8")
        return
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("写入 docx 需要安装 python-docx：pip install -r requirements.txt") from exc
        document = Document()
        for line in text.splitlines():
            document.add_paragraph(line)
        document.save(str(file_path))
        return
    raise ValueError(f"不支持的输出格式：{suffix}")


def redaction_map_to_json(redaction_map: RedactionMap) -> str:
    return json.dumps(redaction_map.to_dict(), ensure_ascii=False, indent=2)


def redaction_map_from_json(value: str) -> RedactionMap:
    return RedactionMap.from_dict(json.loads(value))


def load_redaction_map(path: str | Path) -> RedactionMap:
    return redaction_map_from_json(Path(path).read_text(encoding="utf-8"))


def save_redaction_map(path: str | Path, redaction_map: RedactionMap) -> None:
    file_path = Path(path)
    file_path.parent.mkdir(parents=True, exist_ok=True)
    file_path.write_text(redaction_map_to_json(redaction_map), encoding="utf-8")


def save_redaction_map_encrypted(path: str | Path, redaction_map: RedactionMap) -> None:
    """加密保存 redaction_map，敏感原始数据不会明文落盘。"""
    from ._crypto import encrypt

    file_path = Path(path)
    file_path.parent.mkdir(parents=True, exist_ok=True)
    json_str = redaction_map_to_json(redaction_map)
    encrypted = encrypt(json_str)
    file_path.write_bytes(encrypted)
    os.chmod(file_path, 0o600)


def load_redaction_map_encrypted(path: str | Path) -> RedactionMap:
    """解密加载 redaction_map。"""
    from ._crypto import decrypt

    data = Path(path).read_bytes()
    return redaction_map_from_json(decrypt(data))


def load_redaction_map_auto(path: str | Path) -> RedactionMap:
    """自动读取明文或加密映射表。"""
    if is_encrypted_map(path):
        return load_redaction_map_encrypted(path)
    return load_redaction_map(path)


def is_encrypted_map(path: str | Path) -> bool:
    """判断文件是否为加密格式（不以 '{' 开头）。"""
    try:
        with open(path, "rb") as f:
            head = f.read(1)
        return head != b"{"
    except Exception:
        return False


def save_redaction_map_auto(path: str | Path, redaction_map: RedactionMap) -> None:
    """自动选择加密或明文保存。

    如果密钥可用（环境变量或已持久化），自动加密保存；
    否则明文保存并打印警告。
    """
    from ._crypto import get_or_create_key

    try:
        get_or_create_key()
        save_redaction_map_encrypted(path, redaction_map)
    except Exception:
        import sys

        print(
            "[legal-redactor] 警告：redaction_map 将明文保存。"
            "设置环境变量 LEGAL_REDACTOR_KEY 以启用加密。",
            file=sys.stderr,
        )
        save_redaction_map(path, redaction_map)
